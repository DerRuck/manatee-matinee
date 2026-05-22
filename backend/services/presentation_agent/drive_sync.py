"""Upload validated PresentationOutlines to Google Drive.

Auth and the low-level upload primitive are reused from
services/research_agent/drive_sync so both agents land in the same Drive,
under the same account, with identical idempotency semantics.

Layout produced (root = DEFAULT_FOLDER_ID):
    {root}/{contact folder}/Presentation Outlines/
        pa_curiosity_{run_id8}.json
        pa_curiosity_{run_id8}.docx
        pa_curiosity_{run_id8}.html

Three formats land per run:
  - .json  — canonical machine-readable outline (consumed by Canva/Figma/python-pptx)
  - .docx  — reviewable outline for staff (slide-by-slide cards)
  - .html  — self-contained brand-styled deck (one file, opens in any browser,
             prints to PDF, can be screenshared as-is)

Usage:
    from services.presentation_agent.drive_sync import upload_outline
    result = upload_outline(outline)
    # result = {"json": {...}, "docx": {...}, "html": {...}}
"""

from __future__ import annotations

import base64
import html as html_lib
import re
from pathlib import Path
from typing import Any

from services.presentation_agent.schema import PresentationOutline
# Reuse research-agent infra: same auth, same upload-or-replace semantics.
from services.research_agent.drive_sync import (
    DOCX_MIME,
    DEFAULT_FOLDER_ID,
    _get_drive_service,
    _upload_or_replace,
)

HTML_MIME = "text/html"


# ---------------------------------------------------------------------------
# Filenames
# ---------------------------------------------------------------------------

def _slug(text: str | None) -> str:
    return re.sub(r"[^a-z0-9]+", "_", (text or "unknown").lower()).strip("_")


def filename_for(outline: PresentationOutline, ext: str) -> str:
    return (
        f"{_slug(outline.municipality_name)}_"
        f"{outline.outline_type_id.lower().replace('-', '_')}_"
        f"v{outline.prompt_version}_{outline.run_id[:8]}.{ext}"
    )


# ---------------------------------------------------------------------------
# DOCX renderer — slide-by-slide cards
#
# Staff opens this to review the outline before any deck gets built. Each
# slide is one card with layout-appropriate content + speaker notes. Same
# python-docx pattern as research_agent renderers, simpler structure since
# the slide types are tighter.
# ---------------------------------------------------------------------------

_LAYOUT_LABELS = {
    "title":              "Title",
    "section_divider":    "Section Divider",
    "three_pillar":       "Three Pillars",
    "bullet":             "Bullet Slide",
    "team_bio":           "Team Bios",
    "data_point":         "Data Point",
    "comparable_project": "Comparable Project",
    "quote":              "Quote",
    "closing":            "Closing",
}


def _add_kv(doc: Any, label: str, value: Any) -> None:
    if value is None or value == "":
        return
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(str(value))


def _render_slide_card(doc: Any, slide: Any) -> None:
    layout = slide.layout
    layout_label = _LAYOUT_LABELS.get(layout, layout)

    title = (
        getattr(slide, "title", None)
        or getattr(slide, "section_title", None)
        or getattr(slide, "headline", None)
        or getattr(slide, "project_name", None)
        or getattr(slide, "call_to_action", None)
        or getattr(slide, "quote", None)
        or layout_label
    )
    doc.add_heading(f"Slide {slide.slide_number}: {title}", level=2)
    _add_kv(doc, "Layout", layout_label)

    if layout == "title":
        _add_kv(doc, "Subtitle", slide.subtitle)
        _add_kv(doc, "Date", slide.date_text)
        _add_kv(doc, "Footer", slide.footer_text)

    elif layout == "section_divider":
        _add_kv(doc, "Section #", slide.section_number)
        _add_kv(doc, "Big idea", slide.big_idea)

    elif layout == "three_pillar":
        for i, pillar in enumerate(slide.pillars, 1):
            doc.add_heading(f"Pillar {i}: {pillar.heading}", level=3)
            doc.add_paragraph(pillar.body)

    elif layout == "bullet":
        _add_kv(doc, "Subtitle", slide.subtitle)
        for bullet in slide.bullets:
            doc.add_paragraph(bullet, style="List Bullet")
        for asset in slide.visual_assets:
            doc.add_paragraph(
                f"[visual: {asset.asset_type}] {asset.description}",
                style="Intense Quote",
            )

    elif layout == "team_bio":
        for member in slide.members:
            doc.add_heading(f"{member.name} — {member.role}", level=3)
            if member.bio_one_liner:
                doc.add_paragraph(member.bio_one_liner)
            _add_kv(doc, "Passion", member.passion)
            _add_kv(doc, "Fact", member.fact)

    elif layout == "data_point":
        _add_kv(doc, "Headline", slide.headline)
        _add_kv(doc, "Plain-English framing", slide.plain_english_framing)
        for src in slide.sources:
            _add_kv(doc, "Source", f"{src.title or src.url} (reliability {src.reliability_score})")

    elif layout == "comparable_project":
        _add_kv(doc, "Project", slide.project_name)
        _add_kv(doc, "Municipality", slide.municipality)
        _add_kv(doc, "Year", slide.year)
        if slide.cost_usd:
            _add_kv(doc, "Cost", f"${slide.cost_usd:,}")
        _add_kv(doc, "Outcome", slide.outcome)
        _add_kv(doc, "Why relevant", slide.why_relevant)

    elif layout == "quote":
        p = doc.add_paragraph()
        p.add_run(f"“{slide.quote}”").italic = True
        _add_kv(doc, "Attribution", slide.attribution)
        _add_kv(doc, "Context", slide.context)

    elif layout == "closing":
        _add_kv(doc, "Call to action", slide.call_to_action)
        _add_kv(doc, "Leave-behind", slide.leave_behind_summary)
        _add_kv(doc, "Contact line", slide.contact_line)

    if slide.speaker_notes:
        doc.add_heading("Speaker notes", level=3)
        p = doc.add_paragraph()
        p.add_run(slide.speaker_notes).italic = True


def render_docx(outline: PresentationOutline) -> bytes:
    """Build a .docx outline doc for human review."""
    import io
    from docx import Document

    from services.branding.docx_styles import (
        add_brand_header, add_meta_line, apply_brand_styles,
    )

    doc = Document()
    apply_brand_styles(doc)
    f = outline.findings

    add_brand_header(
        doc,
        title=f.deck_title,
        subtitle=f.deck_subtitle or f"{outline.outline_type_id} · presentation outline",
    )

    add_meta_line(
        doc,
        generated=outline.generated_at.strftime("%B %d, %Y at %H:%M UTC"),
        confidence=outline.overall_confidence,
        run=outline.run_id[:8],
        slides=len(f.slides),
    )

    doc.add_heading("Meeting context", level=1)
    _add_kv(doc, "Outline type", outline.outline_type_id)
    _add_kv(doc, "Audience", getattr(f, "audience", None))
    _add_kv(doc, "Objective", getattr(f, "meeting_objective", None))
    _add_kv(doc, "Champion", getattr(f, "champion_name", None))
    _add_kv(doc, "Municipality", outline.municipality_name)
    _add_kv(doc, "Project", getattr(f, "project_name", None))
    _add_kv(doc, "Problem-area focus", getattr(f, "problem_area_focus", None))

    if outline.upstream_briefs:
        doc.add_heading("Upstream research briefs", level=1)
        for ref in outline.upstream_briefs:
            line = f"{ref.research_type_id}"
            if ref.run_id:
                line += f" (run {ref.run_id[:8]})"
            if ref.summary:
                line += f" — {ref.summary}"
            doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("Slides", level=1)
    for slide in f.slides:
        _render_slide_card(doc, slide)

    cadence = getattr(f, "communication_cadence", None)
    if cadence:
        doc.add_heading("Communication cadence", level=1)
        doc.add_paragraph(cadence)

    risks = getattr(f, "top_risks", None)
    if risks:
        doc.add_heading("Top risks", level=1)
        for risk in risks:
            doc.add_paragraph(risk, style="List Bullet")

    next_step = getattr(f, "suggested_next_step", None)
    if next_step:
        doc.add_heading("Suggested next step", level=1)
        doc.add_paragraph(next_step)

    if outline.notes:
        doc.add_heading("Notes", level=1)
        doc.add_paragraph(outline.notes)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# HTML renderer — self-contained brand-styled deck (1920×1080 per slide)
#
# Styling pulled from the C-HAWQ Design System (Remix): colors_and_type.css
# + slide-templates.css, inlined as one CSS string below. Layouts the kit
# doesn't ship (bullet, team_bio, data_point, comparable_project, quote)
# get extension templates here. Logos embedded as base64 so the .html file
# is fully self-contained — opens in any browser, prints to PDF cleanly
# with one slide per page.
# ---------------------------------------------------------------------------

_BRANDING_ASSETS = Path(__file__).resolve().parent.parent / "branding" / "assets"


def _load_logo_data_url(filename: str, mime: str) -> str:
    """Embed a logo file as a base64 data URL so the deck stays self-contained."""
    path = _BRANDING_ASSETS / "logos" / filename
    if not path.exists():
        return ""
    return f"data:{mime};base64,{base64.b64encode(path.read_bytes()).decode('ascii')}"


# Loaded once at import. Light-on-blue mark for dark slides; full-color symbol
# for light slides.
_LOGO_ON_BLUE = _load_logo_data_url("logo-secondary-symbol-cropped.png", "image/png")
_LOGO_ON_WHITE = _load_logo_data_url("logo-primary-symbol.jpg", "image/jpeg")


_HTML_LAYOUT_LABELS = {
    "title":              "Cover",
    "section_divider":    "Section",
    "three_pillar":       "Three Pillars",
    "bullet":             "Key Points",
    "team_bio":           "Team",
    "data_point":         "Data Point",
    "comparable_project": "Florida Precedent",
    "quote":              "Quote",
    "closing":            "Closing",
}


def _esc(value: Any) -> str:
    return html_lib.escape("" if value is None else str(value), quote=True)


def _chrome_top(label: str, on_blue: bool) -> str:
    """Top-left logo + top-right eyebrow band on every slide."""
    logo = _LOGO_ON_BLUE if on_blue else _LOGO_ON_WHITE
    wm_style = "" if on_blue else ' style="color:var(--chawq-main-blue)"'
    return (
        '<div class="chrome-top">'
        f'<div class="brand"><img src="{logo}" alt="C-HAWQ" />'
        f'<span class="wm"{wm_style}>C-HAWQ</span></div>'
        f'<div class="eyebrow">{_esc(label)}</div>'
        '</div>'
    )


def _chrome_bot(page_num: int, total: int) -> str:
    return (
        '<div class="chrome-bot">'
        '<span>C-HAWQ</span>'
        f'<span class="page">{page_num:02d} / {total:02d}</span>'
        '</div>'
    )


def _speaker_notes(slide: Any) -> str:
    """Speaker notes collapse into a <details>. Hidden in print."""
    if not getattr(slide, "speaker_notes", None):
        return ""
    return (
        '<details class="speaker-notes">'
        '<summary>Speaker notes</summary>'
        f'<p>{_esc(slide.speaker_notes)}</p>'
        '</details>'
    )


def _slide_html(slide: Any, total: int) -> str:
    layout = slide.layout
    label = _HTML_LAYOUT_LABELS.get(layout, layout)
    n = slide.slide_number

    if layout == "title":
        sub = f'<div class="sub">{_esc(slide.subtitle)}</div>' if slide.subtitle else ""
        meta = f'<div class="meta">{_esc(slide.date_text)}</div>' if slide.date_text else ""
        return (
            f'<section class="slide blue t-title">'
            f'{_chrome_top(label, True)}'
            f'<div class="rule"></div>'
            f'<h1>{_esc(slide.title)}</h1>'
            f'{sub}{meta}'
            f'{_chrome_bot(n, total)}'
            f'{_speaker_notes(slide)}'
            f'</section>'
        )

    if layout == "section_divider":
        num = f"{slide.section_number:02d}" if slide.section_number is not None else "—"
        idea = f'<div class="lead">{_esc(slide.big_idea)}</div>' if slide.big_idea else ""
        return (
            f'<section class="slide blue t-section">'
            f'{_chrome_top(label, True)}'
            f'<div class="num">{num} / {_esc(slide.section_title.upper())}</div>'
            f'<h1>{_esc(slide.section_title)}</h1>'
            f'<div class="rule"></div>'
            f'{idea}'
            f'{_chrome_bot(n, total)}'
            f'{_speaker_notes(slide)}'
            f'</section>'
        )

    if layout == "three_pillar":
        cards = "".join(
            f'<div class="card"><div class="num">{i + 1:02d}</div>'
            f'<h3>{_esc(p.heading)}</h3><p>{_esc(p.body)}</p></div>'
            for i, p in enumerate(slide.pillars)
        )
        return (
            f'<section class="slide sky-tint t-pillars">'
            f'{_chrome_top(label, False)}'
            f'<div class="head"><h2>{_esc(slide.title)}</h2></div>'
            f'<div class="grid">{cards}</div>'
            f'{_chrome_bot(n, total)}'
            f'{_speaker_notes(slide)}'
            f'</section>'
        )

    if layout == "bullet":
        sub = f'<div class="lead">{_esc(slide.subtitle)}</div>' if slide.subtitle else ""
        bullets = "".join(f'<li>{_esc(b)}</li>' for b in slide.bullets)
        assets = "".join(
            f'<p class="visual-hint"><span class="visual-tag">[{_esc(a.asset_type)}]</span> '
            f'{_esc(a.description)}</p>'
            for a in getattr(slide, "visual_assets", []) or []
        )
        return (
            f'<section class="slide t-bullet">'
            f'{_chrome_top(label, False)}'
            f'<div class="lhs"><h2>{_esc(slide.title)}</h2>{sub}</div>'
            f'<div class="rhs"><ul class="bullets">{bullets}</ul>{assets}</div>'
            f'{_chrome_bot(n, total)}'
            f'{_speaker_notes(slide)}'
            f'</section>'
        )

    if layout == "team_bio":
        members = "".join(
            (
                '<div class="member">'
                f'<h3>{_esc(m.name)}</h3>'
                f'<div class="role">{_esc(m.role)}</div>'
                + (f'<p>{_esc(m.bio_one_liner)}</p>' if m.bio_one_liner else "")
                + (
                    f'<p class="kv"><span class="kv-label">Passion</span>{_esc(m.passion)}</p>'
                    if m.passion else ""
                )
                + (
                    f'<p class="kv"><span class="kv-label">Fact</span>{_esc(m.fact)}</p>'
                    if m.fact else ""
                )
                + '</div>'
            )
            for m in slide.members
        )
        return (
            f'<section class="slide t-team">'
            f'{_chrome_top(label, False)}'
            f'<div class="head"><h2>{_esc(slide.title)}</h2></div>'
            f'<div class="grid">{members}</div>'
            f'{_chrome_bot(n, total)}'
            f'{_speaker_notes(slide)}'
            f'</section>'
        )

    if layout == "data_point":
        sources = "".join(
            f'<li><a href="{_esc(str(src.url))}">{_esc(src.title or str(src.url))}</a>'
            f'<span class="reliability">r {src.reliability_score:.2f}</span></li>'
            for src in slide.sources
        )
        return (
            f'<section class="slide t-data">'
            f'{_chrome_top(label, False)}'
            f'<div class="fig">{_esc(slide.headline)}</div>'
            f'<p class="framing">{_esc(slide.plain_english_framing)}</p>'
            f'<ul class="sources">{sources}</ul>'
            f'{_chrome_bot(n, total)}'
            f'{_speaker_notes(slide)}'
            f'</section>'
        )

    if layout == "comparable_project":
        meta = f"{_esc(slide.municipality)} · {slide.year}"
        if slide.cost_usd:
            meta += f" · ${slide.cost_usd:,}"
        return (
            f'<section class="slide t-comparable">'
            f'{_chrome_top(label, False)}'
            f'<div class="head">'
            f'<div class="eyebrow">Florida Precedent</div>'
            f'<h2>{_esc(slide.project_name)}</h2>'
            f'<div class="meta">{meta}</div>'
            f'</div>'
            f'<div class="body">'
            f'<p><span class="kv-label">Outcome</span> {_esc(slide.outcome)}</p>'
            f'<p><span class="kv-label">Why this matters</span> {_esc(slide.why_relevant)}</p>'
            f'</div>'
            f'{_chrome_bot(n, total)}'
            f'{_speaker_notes(slide)}'
            f'</section>'
        )

    if layout == "quote":
        context = f'<div class="context">{_esc(slide.context)}</div>' if slide.context else ""
        return (
            f'<section class="slide t-quote">'
            f'{_chrome_top(label, False)}'
            f'<blockquote class="quote">{_esc(slide.quote)}</blockquote>'
            f'<div class="attribution">— {_esc(slide.attribution)}</div>'
            f'{context}'
            f'{_chrome_bot(n, total)}'
            f'{_speaker_notes(slide)}'
            f'</section>'
        )

    if layout == "closing":
        leave = (
            f'<p class="leave">{_esc(slide.leave_behind_summary)}</p>'
            if slide.leave_behind_summary else ""
        )
        return (
            f'<section class="slide blue t-close">'
            f'{_chrome_top(label, True)}'
            f'<div class="rule"></div>'
            f'<h1>{_esc(slide.call_to_action)}</h1>'
            f'{leave}'
            f'<div class="sig">{_esc(slide.contact_line)}</div>'
            f'{_chrome_bot(n, total)}'
            f'{_speaker_notes(slide)}'
            f'</section>'
        )

    # Unknown layout — render a minimal fallback so the deck still produces.
    return (
        f'<section class="slide t-bullet">'
        f'{_chrome_top(label, False)}'
        f'<div class="lhs"><h2>{_esc(layout)}</h2></div>'
        f'{_chrome_bot(n, total)}'
        f'{_speaker_notes(slide)}'
        f'</section>'
    )


_HTML_CSS = """
@import url('https://fonts.googleapis.com/css2?family=Jost:wght@400;500;600;700;800;900&family=Source+Sans+3:ital,wght@0,300;0,400;0,500;0,600;0,700;1,400;1,600&family=Source+Code+Pro:wght@400;500;600&display=swap');

/* ============================================================
   C-HAWQ Design System — verbatim tokens from colors_and_type.css
   ============================================================ */
:root {
  --chawq-main-blue:       #1f396d;
  --chawq-green:           #3f886c;
  --chawq-sky:             #48c5e3;

  --chawq-main-blue-900:   #142749;
  --chawq-main-blue-800:   #1a3160;
  --chawq-main-blue-700:   #1f396d;
  --chawq-main-blue-600:   #2a4a86;
  --chawq-main-blue-500:   #3a5ea0;
  --chawq-main-blue-300:   #7a90bb;
  --chawq-main-blue-100:   #d7dfee;
  --chawq-main-blue-050:   #eef2f9;

  --chawq-green-900:       #255241;
  --chawq-green-700:       #3f886c;
  --chawq-green-500:       #5ba988;
  --chawq-green-100:       #deece5;
  --chawq-green-050:       #eff6f2;

  --chawq-sky-900:         #1d7c94;
  --chawq-sky-700:         #2aa3c2;
  --chawq-sky-500:         #48c5e3;
  --chawq-sky-300:         #8edcee;
  --chawq-sky-100:         #d0eff7;
  --chawq-sky-050:         #ebf8fc;

  --chawq-ink:             #0e1a33;
  --chawq-fg-1:            #152545;
  --chawq-fg-2:            #3c4a66;
  --chawq-fg-3:            #65728c;
  --chawq-fg-muted:        #8a94a9;
  --chawq-line:            #d9dde6;
  --chawq-line-soft:       #e8ebf1;
  --chawq-surface-2:       #f2f4f8;
  --chawq-surface-1:       #f7f8fb;
  --chawq-paper:           #ffffff;

  --chawq-on-blue-1:       #ffffff;
  --chawq-on-blue-2:       #d7dfee;
  --chawq-on-blue-3:       #9aaacb;

  --font-display: 'Jost', 'Futura PT', 'Futura', 'Century Gothic', 'Avenir Next', sans-serif;
  --font-body:    'Source Sans 3', 'Source Sans Pro', 'Open Sans', 'Inter', system-ui, sans-serif;
  --font-mono:    'Source Code Pro', 'Roboto Mono', ui-monospace, Menlo, monospace;
}

* { box-sizing: border-box; }

html { -webkit-font-smoothing: antialiased; text-rendering: optimizeLegibility; }

body {
  margin: 0;
  padding: 24px 0;
  background: var(--chawq-surface-1);
  font-family: var(--font-body);
  color: var(--chawq-fg-1);
  line-height: 1.55;
}

.deck {
  display: flex;
  flex-direction: column;
  align-items: center;
  gap: 24px;
}

/* ============================================================
   Slide base — verbatim from slide-templates.css (1920×1080)
   ============================================================ */
.slide {
  width: 1920px;
  height: 1080px;
  position: relative;
  overflow: hidden;
  font-family: var(--font-body);
  color: var(--chawq-fg-1);
  background: #fff;
  box-shadow: 0 18px 40px -10px rgba(20, 39, 73, 0.22);
}
.slide.blue { background: var(--chawq-main-blue); color: #fff; }
.slide.blue h1, .slide.blue h2, .slide.blue h3 { color: #fff; }
.slide.blue .eyebrow { color: var(--chawq-sky); }
.slide.sky-tint { background: linear-gradient(180deg, var(--chawq-sky-050) 0%, #fff 100%); }

/* Shared chrome */
.chrome-top {
  position: absolute; top: 60px; left: 96px; right: 96px;
  display: flex; justify-content: space-between; align-items: center;
}
.chrome-top .brand { display: flex; align-items: center; gap: 14px; }
.chrome-top .brand img { width: 56px; height: 56px; object-fit: contain; }
.chrome-top .brand .wm {
  font-family: var(--font-display); font-weight: 800;
  font-size: 24px; letter-spacing: 0.02em; color: #fff;
}
.chrome-top .eyebrow {
  font-family: var(--font-display); font-size: 18px; font-weight: 700;
  letter-spacing: 0.14em; text-transform: uppercase;
  color: var(--chawq-sky-900);
}
.slide.blue .chrome-top .eyebrow { color: var(--chawq-sky); }
.chrome-bot {
  position: absolute; bottom: 56px; left: 96px; right: 96px;
  display: flex; justify-content: space-between; align-items: center;
  font-family: var(--font-mono); font-size: 16px; color: var(--chawq-fg-muted);
}
.slide.blue .chrome-bot { color: var(--chawq-on-blue-3); }
.chrome-bot .page { color: inherit; }

/* ---------- t-title (kit) ---------- */
.t-title { display: flex; flex-direction: column; justify-content: flex-end; padding: 140px 120px 200px; }
.t-title .rule { height: 6px; width: 120px; background: var(--chawq-sky); margin-bottom: 40px; }
.t-title h1 { font-family: var(--font-display); font-size: 120px; line-height: 0.94; font-weight: 900; letter-spacing: -0.015em; margin: 0 0 28px; max-width: 1500px; color: #fff; }
.t-title .sub { font-family: var(--font-body); font-size: 32px; color: var(--chawq-on-blue-2); max-width: 1200px; line-height: 1.35; }
.t-title .meta { font-family: var(--font-mono); font-size: 20px; color: var(--chawq-sky); margin-top: 60px; letter-spacing: 0.04em; }

/* ---------- t-section (kit) ---------- */
.t-section { display: flex; flex-direction: column; justify-content: center; align-items: flex-start; padding: 0 160px; }
.t-section .num { font-family: var(--font-mono); font-size: 24px; color: var(--chawq-sky); letter-spacing: 0.14em; margin-bottom: 28px; }
.t-section h1 { font-family: var(--font-display); font-size: 180px; line-height: 0.92; font-weight: 900; letter-spacing: -0.02em; margin: 0; max-width: 1500px; color: #fff; }
.t-section .rule { height: 4px; width: 180px; background: var(--chawq-sky); margin-top: 56px; }
.t-section .lead { font-family: var(--font-body); font-size: 28px; color: var(--chawq-on-blue-2); max-width: 1200px; line-height: 1.4; margin-top: 32px; }

/* ---------- t-pillars (kit) ---------- */
.t-pillars { padding: 170px 96px 140px; }
.t-pillars .head { margin-bottom: 64px; }
.t-pillars h2 { font-family: var(--font-display); font-size: 72px; line-height: 1.02; margin: 0; font-weight: 800; color: var(--chawq-main-blue); max-width: 1500px; }
.t-pillars .grid { display: grid; grid-template-columns: repeat(3, 1fr); gap: 36px; }
.t-pillars .card { border: 1.5px solid var(--chawq-line); border-radius: 18px; padding: 48px 44px; background: #fff; }
.t-pillars .card .num { font-family: var(--font-mono); font-size: 16px; color: var(--chawq-sky-900); font-weight: 600; letter-spacing: 0.08em; margin-bottom: 20px; }
.t-pillars .card h3 { font-family: var(--font-display); font-size: 36px; line-height: 1.1; margin: 0 0 18px; font-weight: 800; color: var(--chawq-main-blue); }
.t-pillars .card p { font-size: 22px; line-height: 1.5; color: var(--chawq-fg-2); margin: 0; }

/* ---------- t-close (kit) ---------- */
.t-close { padding: 0 160px; display: flex; flex-direction: column; justify-content: center; align-items: flex-start; }
.t-close .rule { height: 6px; width: 120px; background: var(--chawq-sky); margin-bottom: 40px; }
.t-close h1 { font-family: var(--font-display); font-size: 96px; line-height: 1; font-weight: 900; margin: 0 0 40px; color: #fff; max-width: 1500px; }
.t-close .leave { font-family: var(--font-body); font-size: 28px; color: var(--chawq-on-blue-2); max-width: 1200px; line-height: 1.4; margin: 0 0 56px; }
.t-close .sig { font-family: var(--font-mono); font-size: 22px; color: var(--chawq-sky); letter-spacing: 0.04em; }

/* ============================================================
   Extension templates — for schema layouts the kit doesn't ship
   ============================================================ */

/* ---------- t-bullet (headline left, bullets right) ---------- */
.t-bullet { padding: 170px 120px 140px; display: grid; grid-template-columns: 1fr 1fr; gap: 100px; align-content: start; }
.t-bullet .lhs h2 { font-family: var(--font-display); font-size: 76px; line-height: 1.02; margin: 0; font-weight: 800; color: var(--chawq-main-blue); max-width: 820px; }
.t-bullet .lhs .lead { font-size: 24px; color: var(--chawq-fg-2); margin-top: 24px; max-width: 700px; line-height: 1.45; }
.t-bullet .rhs ul.bullets { margin: 0; padding-left: 32px; font-size: 28px; line-height: 1.45; color: var(--chawq-fg-2); }
.t-bullet .rhs ul.bullets li { margin-bottom: 22px; }
.t-bullet .rhs ul.bullets li::marker { color: var(--chawq-sky); font-weight: 700; }
.t-bullet .visual-hint { font-family: var(--font-mono); font-size: 14px; color: var(--chawq-fg-muted); font-style: italic; margin-top: 32px; }
.t-bullet .visual-tag { font-family: var(--font-mono); color: var(--chawq-sky-900); font-style: normal; margin-right: 6px; }

/* ---------- t-team (grid of member cards) ---------- */
.t-team { padding: 170px 96px 140px; }
.t-team .head { margin-bottom: 56px; }
.t-team h2 { font-family: var(--font-display); font-size: 72px; line-height: 1.02; margin: 0; font-weight: 800; color: var(--chawq-main-blue); }
.t-team .grid { display: grid; grid-template-columns: repeat(3, 1fr); gap: 32px; }
.t-team .member { border: 1.5px solid var(--chawq-line); border-radius: 14px; padding: 36px 32px; background: #fff; }
.t-team .member h3 { font-family: var(--font-display); font-size: 30px; line-height: 1.1; margin: 0 0 6px; font-weight: 800; color: var(--chawq-main-blue); }
.t-team .member .role { font-family: var(--font-mono); font-size: 14px; letter-spacing: 0.08em; text-transform: uppercase; color: var(--chawq-sky-900); margin-bottom: 18px; }
.t-team .member p { font-size: 18px; line-height: 1.5; color: var(--chawq-fg-2); margin: 0 0 10px; }
.t-team .member .kv-label { font-family: var(--font-display); font-size: 12px; font-weight: 700; letter-spacing: 0.14em; text-transform: uppercase; color: var(--chawq-sky-900); display: block; margin-bottom: 4px; }

/* ---------- t-data (single big stat + sources) ---------- */
.t-data { padding: 170px 120px 140px; display: flex; flex-direction: column; justify-content: center; }
.t-data .fig { font-family: var(--font-display); font-weight: 900; font-size: 120px; line-height: 1; letter-spacing: -0.02em; color: var(--chawq-main-blue); margin: 0 0 40px; max-width: 1600px; }
.t-data .framing { font-family: var(--font-body); font-size: 32px; line-height: 1.4; color: var(--chawq-fg-2); margin: 0 0 56px; max-width: 1400px; }
.t-data ul.sources { list-style: none; margin: 0; padding: 0; }
.t-data ul.sources li { font-family: var(--font-mono); font-size: 16px; color: var(--chawq-fg-3); padding: 8px 0; border-top: 1px solid var(--chawq-line); display: flex; justify-content: space-between; align-items: center; gap: 16px; }
.t-data ul.sources a { color: var(--chawq-main-blue-600); text-decoration: none; }
.t-data ul.sources .reliability { font-family: var(--font-mono); color: var(--chawq-sky-900); }

/* ---------- t-comparable (Florida precedent card) ---------- */
.t-comparable { padding: 170px 120px 140px; display: grid; grid-template-rows: auto 1fr; gap: 56px; }
.t-comparable .head .eyebrow { font-family: var(--font-display); font-size: 18px; font-weight: 700; letter-spacing: 0.14em; text-transform: uppercase; color: var(--chawq-sky-900); margin-bottom: 18px; }
.t-comparable h2 { font-family: var(--font-display); font-size: 76px; line-height: 1; margin: 0; font-weight: 800; color: var(--chawq-main-blue); max-width: 1500px; }
.t-comparable .head .meta { font-family: var(--font-mono); font-size: 22px; color: var(--chawq-fg-3); letter-spacing: 0.04em; margin-top: 24px; }
.t-comparable .body p { font-size: 28px; line-height: 1.45; color: var(--chawq-fg-2); margin: 0 0 28px; max-width: 1500px; }
.t-comparable .body .kv-label { font-family: var(--font-display); font-size: 14px; font-weight: 700; letter-spacing: 0.14em; text-transform: uppercase; color: var(--chawq-sky-900); display: block; margin-bottom: 8px; }

/* ---------- t-quote (large pull quote) ---------- */
.t-quote { padding: 0 160px; display: flex; flex-direction: column; justify-content: center; align-items: flex-start; }
.t-quote .quote { font-family: var(--font-display); font-weight: 700; font-size: 64px; line-height: 1.15; color: var(--chawq-main-blue); border-left: 6px solid var(--chawq-sky); padding-left: 48px; margin: 0 0 48px; max-width: 1500px; }
.t-quote .attribution { font-family: var(--font-display); font-size: 24px; font-weight: 700; letter-spacing: 0.04em; color: var(--chawq-fg-2); margin-bottom: 12px; }
.t-quote .context { font-family: var(--font-mono); font-size: 18px; color: var(--chawq-fg-muted); }

/* ============================================================
   Speaker notes — visible on screen, hidden in print
   ============================================================ */
.speaker-notes {
  width: 1920px; margin: 0 auto;
  padding: 16px 96px;
  background: var(--chawq-surface-2);
  border-top: 1px solid var(--chawq-line);
  font-size: 18px;
  color: var(--chawq-fg-2);
}
.speaker-notes summary {
  cursor: pointer; font-family: var(--font-display); font-weight: 700;
  letter-spacing: 0.06em; text-transform: uppercase;
  color: var(--chawq-main-blue); font-size: 14px;
}
.speaker-notes p {
  margin: 12px 0 0; padding-left: 16px;
  border-left: 3px solid var(--chawq-sky);
  font-style: italic;
}

/* ============================================================
   Appendix (cadence / risks / next step) — paper card after deck
   ============================================================ */
.appendix {
  width: 1920px; margin: 24px auto 0; padding: 56px 96px;
  background: var(--chawq-paper);
  border-top: 6px solid var(--chawq-green);
}
.appendix-block { margin-bottom: 40px; }
.appendix-block h2 {
  font-family: var(--font-display); font-weight: 800;
  font-size: 32px; color: var(--chawq-main-blue); margin: 0 0 16px;
}
.appendix-block p, .appendix-block li {
  font-family: var(--font-body); font-size: 22px; line-height: 1.55;
  color: var(--chawq-fg-2); margin: 6px 0;
}

/* ============================================================
   Print — one slide per page at native 1920×1080
   ============================================================ */
@page { size: 1920px 1080px; margin: 0; }
@media print {
  body { background: #fff; padding: 0; }
  .deck { gap: 0; }
  .slide { box-shadow: none; page-break-after: always; break-after: page; }
  .speaker-notes { display: none; }
  .appendix { page-break-before: always; }
}
"""


def render_html(outline: PresentationOutline) -> bytes:
    """Build a self-contained brand-styled HTML deck.

    One file, opens in any browser, prints to PDF cleanly. Each slide is a
    16:9 styled section. Speaker notes collapse into a <details> per slide.
    """
    f = outline.findings
    total = len(f.slides)

    slides_html = "\n".join(_slide_html(s, total) for s in f.slides)

    head = (
        '<!DOCTYPE html>\n'
        '<html lang="en">\n'
        '<head>\n'
        '<meta charset="utf-8">\n'
        f'<title>{_esc(f.deck_title)}</title>\n'
        '<meta name="viewport" content="width=device-width, initial-scale=1">\n'
        f'<style>{_HTML_CSS}</style>\n'
        '</head>\n'
        '<body>\n'
    )

    body = '<main class="deck">\n' + slides_html + '\n</main>\n'

    appendix_bits: list[str] = []
    cadence = getattr(f, "communication_cadence", None)
    if cadence:
        appendix_bits.append(
            '<section class="appendix-block">'
            '<h2>Communication cadence</h2>'
            f'<p>{_esc(cadence)}</p>'
            '</section>'
        )
    risks = getattr(f, "top_risks", None) or []
    if risks:
        risk_items = "".join(f"<li>{_esc(r)}</li>" for r in risks)
        appendix_bits.append(
            '<section class="appendix-block">'
            '<h2>Top risks</h2>'
            f'<ul>{risk_items}</ul>'
            '</section>'
        )
    next_step = getattr(f, "suggested_next_step", None)
    if next_step:
        appendix_bits.append(
            '<section class="appendix-block">'
            '<h2>Suggested next step</h2>'
            f'<p>{_esc(next_step)}</p>'
            '</section>'
        )

    appendix = (
        '<aside class="appendix">\n' + "\n".join(appendix_bits) + '\n</aside>\n'
        if appendix_bits else ""
    )

    foot = '</body>\n</html>\n'

    return (head + body + appendix + foot).encode("utf-8")


# ---------------------------------------------------------------------------
# Upload
# ---------------------------------------------------------------------------

def upload_outline(
    outline: PresentationOutline,
    folder_id: str = DEFAULT_FOLDER_ID,
) -> dict[str, dict]:
    """Upload outline as JSON + Word doc + HTML deck to Drive.

    Returns file metadata per format. Files land in
    {folder_id}/{contact folder}/Presentation Outlines/. The contact folder
    and "Presentation Outlines" subfolder are created on first use and
    reused on re-runs.
    """
    from services.drive.folders import ensure_subfolder, resolve_contact_folder_name

    service = _get_drive_service()

    contact_folder_id = ensure_subfolder(
        service, folder_id, resolve_contact_folder_name(outline),
    )
    target_folder_id = ensure_subfolder(
        service, contact_folder_id, "Presentation Outlines",
    )

    json_bytes = outline.model_dump_json(indent=2).encode("utf-8")
    docx_bytes = render_docx(outline)
    html_bytes = render_html(outline)

    return {
        "json": _upload_or_replace(
            service, filename_for(outline, "json"),
            json_bytes, "application/json", target_folder_id,
        ),
        "docx": _upload_or_replace(
            service, filename_for(outline, "docx"),
            docx_bytes, DOCX_MIME, target_folder_id,
        ),
        "html": _upload_or_replace(
            service, filename_for(outline, "html"),
            html_bytes, HTML_MIME, target_folder_id,
        ),
    }
