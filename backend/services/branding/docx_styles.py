"""
C-HAWQ brand styling for python-docx Documents.

Source of truth: Marketing Guide. The print fonts there (Futura PT,
Source Sans 3, Source Code Pro) aren't available in Google Workspace,
so the guide specifies digital substitutions — Poppins, Open Sans,
Roboto Mono. Staff opens these docs in Google Docs, so we set the
substitutions directly.

Usage in a renderer:

    from docx import Document
    from services.branding.docx_styles import (
        apply_brand_styles, add_brand_header, add_meta_line, confidence_run,
    )

    doc = Document()
    apply_brand_styles(doc)
    add_brand_header(doc, title="Lobbyist Registration Check", subtitle="Cedar Key, FL")
    add_meta_line(doc, generated="May 19, 2026", run="8f83f528", model="claude-sonnet-4-6")
    # ... normal doc.add_heading / doc.add_paragraph calls inherit brand fonts ...

The module touches only python-docx — no rendering logic, no Pydantic
schemas. That's why both research and presentation renderers can adopt
it without coupling to either domain.
"""

from __future__ import annotations

from typing import Any

from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Pt, RGBColor

# ---------------------------------------------------------------------------
# Brand palette — Marketing Guide §03
# ---------------------------------------------------------------------------

MAIN_BLUE = RGBColor(0x1F, 0x39, 0x6D)   # primary — headings, structure
GREEN     = RGBColor(0x3F, 0x88, 0x6C)   # success — high-confidence signals
SKY       = RGBColor(0x48, 0xC5, 0xE3)   # accent — moderate confidence, highlights
INK       = RGBColor(0x14, 0x1F, 0x36)   # body text
SUBTLE    = RGBColor(0x55, 0x66, 0x88)   # captions, meta, subtitles
WARNING   = RGBColor(0xC8, 0x4A, 0x3A)   # low-confidence escalation
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)

# ---------------------------------------------------------------------------
# Typography — Marketing Guide §04 (Digital / Google substitutions)
# ---------------------------------------------------------------------------

FONT_DISPLAY = "Poppins"        # Futura PT substitute — titles + headings
FONT_BODY    = "Open Sans"      # Source Sans 3 substitute — paragraphs
FONT_MONO    = "Roboto Mono"    # Source Code Pro substitute — data, IDs


# ---------------------------------------------------------------------------
# Style application
# ---------------------------------------------------------------------------

def _set_style(
    doc: Any,
    style_name: str,
    *,
    font: str,
    size_pt: int,
    color: RGBColor,
    bold: bool = False,
    italic: bool = False,
    space_after_pt: int | None = None,
) -> None:
    """Mutate an existing built-in style to use brand specs.

    python-docx exposes every doc's built-in styles via doc.styles[name].
    Modifying them here means any subsequent doc.add_heading() / paragraph
    inherits the brand look — renderer code doesn't change.
    """
    style = doc.styles[style_name]
    f = style.font
    f.name = font
    # East-Asian font name (rFonts.eastAsia) — Word uses this when the document
    # locale falls back; setting it keeps Google Docs from re-substituting.
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font}" w:hAnsi="{font}" w:cs="{font}"/>')
        rpr.append(rfonts)
    else:
        rfonts.set(qn("w:ascii"), font)
        rfonts.set(qn("w:hAnsi"), font)
        rfonts.set(qn("w:cs"), font)

    f.size = Pt(size_pt)
    f.color.rgb = color
    f.bold = bold
    f.italic = italic
    if space_after_pt is not None:
        style.paragraph_format.space_after = Pt(space_after_pt)


def apply_brand_styles(doc: Any) -> None:
    """Apply C-HAWQ brand styling to every built-in style in this Document.

    Call once, right after `doc = Document()`. After this, every call to
    doc.add_heading / doc.add_paragraph / doc.add_paragraph(style="List Bullet")
    picks up brand fonts and colors. No other renderer changes needed.
    """
    # Display fonts — headings sit in navy, sized for hierarchy
    _set_style(doc, "Title",     font=FONT_DISPLAY, size_pt=26, color=MAIN_BLUE, bold=True,  space_after_pt=6)
    _set_style(doc, "Heading 1", font=FONT_DISPLAY, size_pt=18, color=MAIN_BLUE, bold=True,  space_after_pt=4)
    _set_style(doc, "Heading 2", font=FONT_DISPLAY, size_pt=14, color=MAIN_BLUE, bold=True,  space_after_pt=2)
    _set_style(doc, "Heading 3", font=FONT_DISPLAY, size_pt=12, color=MAIN_BLUE, bold=True,  space_after_pt=2)

    # Body fonts — ink-dark, comfortable read width
    _set_style(doc, "Normal",      font=FONT_BODY, size_pt=11, color=INK,    space_after_pt=4)
    _set_style(doc, "List Bullet", font=FONT_BODY, size_pt=11, color=INK,    space_after_pt=2)
    _set_style(doc, "List Number", font=FONT_BODY, size_pt=11, color=INK,    space_after_pt=2)

    # Subtler styles
    _set_style(doc, "Intense Quote", font=FONT_BODY, size_pt=10, color=SUBTLE, italic=True, space_after_pt=4)
    _set_style(doc, "Subtitle",      font=FONT_BODY, size_pt=12, color=SUBTLE, italic=True, space_after_pt=6)


# ---------------------------------------------------------------------------
# Branded header + meta line
# ---------------------------------------------------------------------------

def _shaded_cell(cell, hex_color: str) -> None:
    """Fill a table cell with a solid color (python-docx has no first-class API for this)."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_brand_header(doc: Any, title: str, subtitle: str | None = None) -> None:
    """Navy brand band + document title + optional subtitle.

    Replaces the conventional `doc.add_heading(title, level=0)` call so
    every C-HAWQ artifact opens with the same visual signature.
    """
    bar = doc.add_table(rows=1, cols=1)
    bar.autofit = True
    cell = bar.rows[0].cells[0]
    _shaded_cell(cell, "1f396d")
    p = cell.paragraphs[0]
    run = p.add_run("C-HAWQ — Coastal Habitat and Water Quality Initiative")
    run.font.name = FONT_DISPLAY
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = WHITE

    doc.add_heading(title, level=0)

    if subtitle:
        p = doc.add_paragraph()
        r = p.add_run(subtitle)
        r.font.name = FONT_BODY
        r.font.size = Pt(11)
        r.font.color.rgb = SUBTLE
        r.italic = True


def add_meta_line(doc: Any, **kv: Any) -> None:
    """Compact one-line meta strip: 'Generated: ... · Run: ... · Confidence: 0.82'.

    Labels render in display font / navy; values in mono / ink to signal
    'this is data, not prose' per Marketing Guide §04 Rule 3.
    """
    p = doc.add_paragraph()
    for i, (key, value) in enumerate(kv.items()):
        if i > 0:
            sep = p.add_run("   ·   ")
            sep.font.name = FONT_BODY
            sep.font.size = Pt(9)
            sep.font.color.rgb = SUBTLE

        label = p.add_run(f"{_humanize(key)}: ")
        label.font.name = FONT_DISPLAY
        label.font.size = Pt(9)
        label.font.bold = True
        label.font.color.rgb = MAIN_BLUE

        val_str = str(value)
        is_confidence = key.lower().startswith("confidence") and isinstance(value, (int, float))
        val_run = p.add_run(val_str)
        val_run.font.name = FONT_MONO
        val_run.font.size = Pt(9)
        val_run.font.color.rgb = _confidence_color(float(value)) if is_confidence else INK
        if is_confidence:
            val_run.font.bold = True


def _humanize(key: str) -> str:
    return key.replace("_", " ").title()


def _confidence_color(score: float) -> RGBColor:
    """Tier color signaling action threshold, not adjectives.

    >=0.80 green (proceed), 0.50-0.79 sky (moderate, verify), <0.50 warning.
    """
    if score >= 0.80:
        return GREEN
    if score >= 0.50:
        return SKY
    return WARNING


# ---------------------------------------------------------------------------
# Inline helpers a renderer can call when it needs a styled run
# ---------------------------------------------------------------------------

def confidence_run(paragraph: Any, score: float) -> None:
    """Append a tiered confidence chip to an existing paragraph.

    Example: 'Confidence: ' + confidence_run(p, 0.82)  ->  '0.82 (HIGH)' in green mono.
    """
    label = "HIGH" if score >= 0.80 else "MODERATE" if score >= 0.50 else "REVIEW"
    run = paragraph.add_run(f"{score:.2f} ({label})")
    run.font.name = FONT_MONO
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = _confidence_color(score)


def mono_run(paragraph: Any, text: str, *, size: int = 10, color: RGBColor = INK) -> None:
    """Append a monospace run — for run IDs, statistics, citation counts."""
    run = paragraph.add_run(text)
    run.font.name = FONT_MONO
    run.font.size = Pt(size)
    run.font.color.rgb = color


def section_rule(doc: Any) -> None:
    """A thin navy rule paragraph — visual section break."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="6" w:space="1" w:color="1f396d"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
