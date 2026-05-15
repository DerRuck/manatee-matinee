"""Upload validated ResearchBriefs to Google Drive as Word documents.

Auth precedence: DRIVE_SA_EMAIL (impersonate via ADC) > DRIVE_SA_KEY (file) > ADC.

Usage:
    from services.research_agent.drive_sync import upload_brief
    result = upload_brief(brief, folder_id="1L-zcN4jA83EfsrRyei_ewbKOEKMKz-lC")
    # result = {"docx": {...Drive file metadata...}}
"""
from __future__ import annotations

import io
import os
import re
import sys
from typing import Any

from services.research_agent.schema import ResearchBrief

DEFAULT_FOLDER_ID = "1L-zcN4jA83EfsrRyei_ewbKOEKMKz-lC"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

_SKIP_ALWAYS = {"research_type"}

# Metadata fields that add noise inside nested models — hidden from body
_SKIP_NESTED = {"sources", "confidence", "reliability_score", "fetched_at"}

# Chars above which a string gets its own heading + paragraph vs. "Label: value" inline
_PROSE_THRESHOLD = 120

# Human-readable names for every research type
_TYPE_NAMES: dict[str, str] = {
    "LOBBY-1": "Lobbyist Registration Check",
    "PW-1":    "Conference Attendee Research",
    "PW-3":    "Municipality Background Research",
    "S1-2":    "LinkedIn Connection Prep",
    "S1-4":    "Contact Background Research",
    "S3-PREP": "Pre-Meeting Research Package",
    "S3-3":    "Commission Meeting Prep",
    "S4-DECK": "Presentation Deck Research",
    "S5-1":    "Internal Presentation Prep",
    "S5-2":    "Post-Meeting Debrief",
    "S6-1":    "Grant Opportunity Research",
    "S6-2":    "Project Narrative Draft",
    "S6-3":    "Commission Presentation Prep",
    "S7-1":    "Post-Event Debrief",
    "S8-1":    "Political Landscape Mapping",
    "S8-2":    "Community Support Letter",
    "S8-3":    "Politician-Friendly Briefing",
    "S9-1":    "Project Kickoff Deck",
    "S9-2":    "Media & Reporter Research",
    "S9-3":    "Grant Compliance Checklist",
    "S9-4":    "P3 Proposal & RFP Draft",
    "S9-5":     "Partnership Agreement Summary",
    "S10-1":    "Project Case Study",
    "S10-2":    "Referral Outreach Research",
    "S4-LETTER": "Champion Briefing Letter",
    "S7-PLAN":  "Community Event Plan",
}

# Override auto-generated Title Case labels with friendlier names
_FIELD_LABELS: dict[str, str] = {
    "connection_request_message":        "Connection Request",
    "follow_up_message_if_no_reply":     "Follow-Up Message (No Reply)",
    "common_ground_hooks":               "Common Ground",
    "environmental_signals":             "Environmental Signals",
    "three_week_action_plan":            "3-Week Action Plan",
    "per_commissioner_briefs":           "Commissioner Notes",
    "agenda_summary_plain_language":     "Agenda Summary",
    "agenda_items_affecting_chawq":      "Items Affecting C-HAWQ",
    "anticipated_objections":            "Anticipated Objections",
    "top_objections":                    "Top Objections",
    "five_slide_outline":                "5-Slide Outline",
    "political_objection_responses":     "Objection Responses",
    "local_angle":                       "Local Angle",
    "leave_behind":                      "Leave-Behind",
    "suggested_ask":                     "Suggested Ask",
    "why_chawq_answer":                  "Why C-HAWQ?",
    "six_month_checkin_email":           "6-Month Check-In Email",
    "warm_intro_email":                  "Warm Intro Email",
    "thank_you_email":                   "Thank-You Email",
    "recommended_next_action":           "Recommended Next Action",
    "remaining_blockers_before_step_6":  "Remaining Blockers",
    "best_chawq_talking_point":          "Best C-HAWQ Talking Point",
    "program_requirements_summary":      "Program Requirements",
    "compliance_checklist":              "Compliance Checklist",
    "day_one_records":                   "Day-One Records to Establish",
    "common_pitfalls":                   "Common Pitfalls",
    "proposal_sections":                 "Proposal Sections",
    "complexity_factors":                "Complexity Factors",
    "gc_evaluation_criteria":            "GC Evaluation Criteria",
    "procurement_conflicts":             "Procurement Conflicts",
    "plain_language_summary":            "Plain-Language Summary",
    "risk_flags":                        "Risk Flags",
    "attorney_questions":                "Questions for Your Attorney",
    "one_sided_terms":                   "One-Sided Terms",
    "complexity_factors_covered":        "Complexity Factors Coverage",
    "leave_behind_case_study":           "Leave-Behind Case Study",
    "website_highlight":                 "Website Highlight",
    "press_pitch":                       "Press Pitch",
    "pull_quotes":                       "Pull Quotes",
    "before_after_narrative":            "Before / After Narrative",
    "before_after_story":                "Before / After Story",
    "contact_overview":                  "Contact Overview",
    "recent_projects_or_proposals":      "Recent Projects & Proposals",
    "follow_up_email_starters":          "Follow-Up Email Starters",
    "next_48h_actions":                  "Next 48-Hour Actions",
    "next_event_improvements":           "Improvements for Next Event",
    "stakeholder_reads":                 "Stakeholder Reads",
    "key_concerns":                      "Key Concerns",
    "opening_narrative":                 "Opening Narrative",
    "department_interests":              "Department Interests",
    "closing_statement":                 "Closing Statement",
    "recommended_data_visuals":          "Recommended Data Visuals",
    "opening_script":                    "Opening Script",
    "things_not_to_say":                 "What Not to Say",
    "champion_profile":                  "Champion Profile",
    "municipality_overview":             "Municipality Overview",
    "project_intelligence":              "Project Intelligence",
    "funding_landscape_summary":         "Funding Landscape",
    "political_landscape_summary":       "Political Landscape",
    "tailored_discovery_questions":      "Discovery Questions",
    "custom_chawq_intro":                "C-HAWQ Intro",
    "grants":                            "Grant Opportunities",
    "risks_and_disqualifiers":           "Risks & Disqualifiers",
    "p3_contractors":                    "P3 Contractors",
    "active_or_stalled_projects":        "Active / Stalled Projects",
    "recent_budget_news":                "Recent Budget News",
    "environmental_issues":              "Environmental Issues",
    "commission_makeup":                 "Commission Makeup",
    "peer_comparables":                  "Peer Comparables",
    "commissioner_profiles":             "Commissioner Profiles",
    "high_leverage_community_voices":    "High-Leverage Community Voices",
    "top_derailers_and_mitigations":     "Top Derailers & Mitigations",
    "satellite_imagery_sources":         "Satellite Imagery Sources",
    "localized_environmental_data":      "Localized Data",
    "comparable_project_examples":       "Comparable Projects",
    "human_scale_statistics":            "Human-Scale Statistics",
    "visual_storytelling_suggestions":   "Visual Storytelling Suggestions",
    "executive_summary":                 "Executive Summary",
    "grant_narrative":                   "Grant Narrative",
    "urgent_statistics":                 "Urgent Statistics",
    "wire_syndication_assessment":       "Wire Syndication Assessment",
    "pitch_emails":                      "Pitch Emails",
    "local_outlets":                     "Local Outlets",
    "regional_environmental_press":      "Regional Environmental Press",
    "municipal_trade_press":             "Municipal Trade Press",
    "beat_reporters":                    "Beat Reporters",
    "open_items_register":               "Open Items Register",
    "registration_required":             "Registration Required",
    "timing_requirement":                "Timing Requirement",
    "reporting_requirements":            "Reporting Requirements",
    "county_covers_municipalities":      "County Covers Municipalities",
    "top_priority_contacts":             "Priority Contacts",
    "public_statements_on_environment":  "Public Statements on Environment",
    "projects_or_initiatives_led":       "Projects & Initiatives Led",
    "municipality_environmental_challenges": "Municipal Environmental Challenges",
    "mutual_connections":                "Mutual Connections",
    "personalized_opening_line":         "Personalized Opening Line",
    "internal_summary":                  "Internal Summary",
    "public_comment_statement":          "Public Comment Statement",
    "support_letter":                    "Support Letter",
    "talking_points_to_emphasize":       "Talking Points to Emphasize",
    "talking_points_to_avoid":           "Talking Points to Avoid",
    "likely_position_on_relevant_items": "Likely Position",
    "special_contract_language_needed":  "Special Contract Language Needed",
    "suggested_alternative_language":    "Suggested Alternative Language",
    "conflict_description":              "Conflict",
    "why_important":                     "Why It Matters",
    "human_framing":                     "Human Framing",
    "speaker_notes":                     "Speaker Notes",
    "bullet_points":                     "Key Points",
    "suggested_opener":                  "Suggested Opener",
    "why_effective_for_this_audience":   "Why Effective for This Audience",
    "why_relevant":                      "Why Relevant",
    "outcome_summary":                   "Outcome",
    "most_likely_objection":             "Most Likely Objection",
    "most_effective_counter":            "Most Effective Counter",
    "leverage_summary":                  "Leverage",
    "suggested_action":                  "Suggested Action",
    "recommended_resolution":            "Recommended Resolution",
    "likely_owner":                      "Likely Owner",
    "key_talking_points":                "Key Talking Points",
    "primary_concerns":                  "Primary Concerns",
    "email_starter":                     "Email Starter",
    "suggested_subject":                 "Suggested Subject",
    "contact_role_or_affiliation":       "Role / Affiliation",
    "deadline_hours":                    "Deadline (hours)",
    "suggested_improvement":             "Suggested Improvement",
    "relevant_beat":                     "Relevant Beat",
    "recent_relevant_byline":            "Recent Byline",
    "target_outlet_or_reporter":         "Target",
    "documentation_needed":              "Documentation Needed",
    "documentation_required":            "Documentation Required",
    "p3_compatible":                     "P3 Compatible",
    "typical_award_usd_min":             "Typical Award (min)",
    "typical_award_usd_max":             "Typical Award (max)",
    "eligibility_summary":               "Eligibility",
    "deadline_or_cycle":                 "Deadline / Cycle",
    "florida_precedents":                "Florida Precedents",
    "prior_environmental_involvement":   "Prior Environmental Involvement",
    "public_statements":                 "Public Statements",
    "background_summary":                "Background",
    "intro_potential":                   "Intro Potential",
    "affects_chawq_project":             "Affects C-HAWQ Project",
    "impact_description":                "Impact",
    "item_number":                       "Item #",
    "environmental_stance":              "Environmental Stance",
    "seat_or_district":                  "Seat / District",
    "waterbody_or_location":             "Waterbody / Location",
    "imagery_provider":                  "Imagery Provider",
    "suggested_view":                    "Suggested View",
    "cost_usd":                          "Cost",
    "award_usd":                         "Award",
    "start_year":                        "Start",
    "end_year":                          "End",
    "tenure_years":                      "Tenure",
    "last_known_update":                 "Last Known Update",
    "issue_type":                        "Issue Type",
    "project_executed":                  "Project Executed",
    "term_or_section":                   "Term / Section",
    "weight_or_priority":                "Priority",
    "commissioner_type_likely":          "Likely Commissioner Type",
    "likely_source":                     "Likely Source",
    "name_or_org":                       "Name / Organization",
    "name_or_role":                      "Name / Role",
}


# ---------------------------------------------------------------------------
# Public helpers
# ---------------------------------------------------------------------------

def _slug(text: str | None) -> str:
    return re.sub(r"[^a-z0-9]+", "_", (text or "unknown").lower()).strip("_")


def filename_for(brief: ResearchBrief, ext: str) -> str:
    return (
        f"{_slug(brief.municipality_name)}_"
        f"{brief.research_type_id.lower().replace('-', '_')}_"
        f"v{brief.prompt_version}_{brief.run_id[:8]}.{ext}"
    )


# ---------------------------------------------------------------------------
# Word document renderer
# ---------------------------------------------------------------------------

def _label(field_name: str) -> str:
    return _FIELD_LABELS.get(field_name) or field_name.replace("_", " ").title()


def _item_heading(obj: Any, index: int) -> str:
    for attr in (
        "name", "outlet_name", "contact_name", "firm_name",
        "section_title", "requirement", "criterion", "factor",
        "question", "objection", "pitfall", "standard_term",
        "term_or_section", "action", "task", "department",
        "slide_number", "week",
        "concern", "observation", "description",
        "issue", "project_name", "metric", "statistic", "name_or_role",
    ):
        val = getattr(obj, attr, None)
        if val is not None:
            return str(val)
    return f"Item {index}"


def _render_email_draft(doc: Any, email: Any) -> None:
    p = doc.add_paragraph()
    p.add_run("Subject: ").bold = True
    p.add_run(email.subject)
    doc.add_paragraph(email.body)


def _render_value(doc: Any, value: Any, level: int) -> None:
    from pydantic import BaseModel
    from services.research_agent.schema import Claim, EmailDraft

    if value is None or value == "" or value == []:
        return

    if isinstance(value, EmailDraft):
        _render_email_draft(doc, value)

    elif isinstance(value, Claim):
        doc.add_paragraph(value.statement, style="List Bullet")

    elif isinstance(value, str):
        doc.add_paragraph(value)

    elif isinstance(value, (int, float)):
        doc.add_paragraph(str(value))

    elif isinstance(value, bool):
        doc.add_paragraph("Yes" if value else "No")

    elif isinstance(value, list):
        from services.research_agent.schema import Claim as _Claim
        if all(isinstance(v, _Claim) for v in value):
            for claim in value:
                doc.add_paragraph(claim.statement, style="List Bullet")
        elif all(isinstance(v, str) for v in value):
            for item in value:
                doc.add_paragraph(item, style="List Bullet")
        else:
            for i, item in enumerate(value, 1):
                if isinstance(item, BaseModel):
                    heading_text = _item_heading(item, i)
                    doc.add_heading(heading_text, min(level, 9))
                    _render_model(doc, item, level + 1, skip_fields=_SKIP_NESTED)
                else:
                    doc.add_paragraph(str(item), style="List Bullet")

    elif isinstance(value, BaseModel):
        _render_model(doc, value, level, skip_fields=_SKIP_NESTED)

    else:
        doc.add_paragraph(str(value))


def _render_model(doc: Any, model: Any, level: int, skip_fields: set) -> None:
    from pydantic import BaseModel
    from services.research_agent.schema import Claim, EmailDraft

    for field_name in model.model_fields:
        if field_name in skip_fields:
            continue
        value = getattr(model, field_name)
        if value is None or value == "" or value == []:
            continue

        field_label = _label(field_name)

        # EmailDraft → email format
        if isinstance(value, EmailDraft):
            doc.add_heading(field_label, min(level, 9))
            _render_email_draft(doc, value)

        # list[Claim] → clean bullet statements, no metadata
        elif isinstance(value, list) and value and all(isinstance(v, Claim) for v in value):
            doc.add_heading(field_label, min(level, 9))
            for claim in value:
                doc.add_paragraph(claim.statement, style="List Bullet")

        elif isinstance(value, str) and len(value) > _PROSE_THRESHOLD:
            doc.add_heading(field_label, min(level, 9))
            doc.add_paragraph(value)

        elif isinstance(value, str):
            p = doc.add_paragraph()
            p.add_run(f"{field_label}: ").bold = True
            p.add_run(value)

        elif isinstance(value, (int, float)):
            p = doc.add_paragraph()
            p.add_run(f"{field_label}: ").bold = True
            p.add_run(str(value))

        elif isinstance(value, bool):
            p = doc.add_paragraph()
            p.add_run(f"{field_label}: ").bold = True
            p.add_run("Yes" if value else "No")

        elif isinstance(value, list) and value:
            doc.add_heading(field_label, min(level, 9))
            _render_value(doc, value, level + 1)

        elif isinstance(value, BaseModel):
            doc.add_heading(field_label, min(level, 9))
            _render_model(doc, value, level + 1, skip_fields=_SKIP_NESTED)


# ---------------------------------------------------------------------------
# Type-specific renderers
# Registered in _CUSTOM_RENDERERS; called instead of the generic _render_model
# for types where layout matters as much as content.
# ---------------------------------------------------------------------------

def _tbl_header(table: Any, *labels: str) -> None:
    """Write bold header cells into the first row of a table."""
    for cell, label in zip(table.rows[0].cells, labels):
        p = cell.paragraphs[0]
        p.clear()
        p.add_run(label).bold = True


def _render_s8_2(doc: Any, findings: Any) -> None:
    """S8-2 — Community Support Letter: letter body + spoken statement."""
    doc.add_heading("Support Letter", 1)
    doc.add_paragraph(
        "Ready to print on organization letterhead. Edit the signature block before sending.",
    ).italic = True
    doc.add_paragraph(findings.support_letter)

    doc.add_heading("Public Comment Statement", 1)
    doc.add_paragraph(
        "2–3 minute spoken adaptation for commission meeting public comment.",
    ).italic = True
    doc.add_paragraph(findings.public_comment_statement)


def _render_s8_3(doc: Any, findings: Any) -> None:
    """S8-3 — Politician-Friendly Briefing: slide cards, objection table, leave-behind."""
    doc.add_heading("5-Slide Outline", 1)
    for slide in findings.five_slide_outline:
        doc.add_heading(f"Slide {slide.slide_number}: {slide.title}", 2)
        for pt in slide.talking_points:
            doc.add_paragraph(pt, style="List Bullet")

    doc.add_heading("Objection Responses", 1)
    for obj in findings.political_objection_responses:
        doc.add_heading(obj.objection, 3)
        doc.add_paragraph(obj.response)

    if findings.local_angle:
        doc.add_heading("Local Angle", 1)
        doc.add_paragraph(findings.local_angle)

    if findings.leave_behind:
        doc.add_heading("Leave-Behind (Plain-Language One-Pager)", 1)
        doc.add_paragraph(findings.leave_behind)

    if findings.suggested_ask:
        doc.add_heading("Suggested Ask", 1)
        doc.add_paragraph(findings.suggested_ask)


def _render_s9_3(doc: Any, findings: Any) -> None:
    """S9-3 — Grant Compliance Checklist: table-based checklist + pitfalls + email."""
    doc.add_heading("Program Requirements", 1)
    doc.add_paragraph(findings.program_requirements_summary)

    doc.add_heading("Compliance Checklist", 1)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    _tbl_header(tbl, "Requirement", "When Due", "Documentation Needed")
    for req in findings.compliance_checklist:
        row = tbl.add_row().cells
        row[0].text = req.requirement
        row[1].text = req.timing
        row[2].text = "\n".join(f"• {d}" for d in req.documentation_needed)

    doc.add_heading("Day-One Records to Establish", 1)
    for item in findings.day_one_records:
        doc.add_paragraph(item, style="List Bullet")

    if findings.common_pitfalls:
        doc.add_heading("Common Pitfalls", 1)
        for pitfall in findings.common_pitfalls:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(pitfall.pitfall).bold = True
            doc.add_paragraph(f"Prevention: {pitfall.prevention}")

    if findings.six_month_checkin_email:
        doc.add_heading("6-Month Check-In Email", 1)
        _render_email_draft(doc, findings.six_month_checkin_email)


def _render_s9_4(doc: Any, findings: Any) -> None:
    """S9-4 — P3 Proposal: real document sections + evaluation criteria table."""
    doc.add_heading("Proposal", 1)
    for section in findings.proposal_sections:
        doc.add_heading(section.section_title, 2)
        doc.add_paragraph(section.content)

    doc.add_heading("Complexity Factors & Required Contract Language", 1)
    for cf in findings.complexity_factors:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(cf.factor + ": ").bold = True
        p.add_run(cf.special_contract_language_needed)

    doc.add_heading("GC Evaluation Criteria", 1)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    _tbl_header(tbl, "Criterion", "Priority", "Rationale")
    for gc in findings.gc_evaluation_criteria:
        row = tbl.add_row().cells
        row[0].text = gc.criterion
        row[1].text = gc.weight_or_priority.upper()
        row[2].text = gc.rationale

    if findings.procurement_conflicts:
        doc.add_heading("Procurement Conflicts", 1)
        for pc in findings.procurement_conflicts:
            doc.add_heading(pc.standard_term, 3)
            p = doc.add_paragraph()
            p.add_run("Conflict: ").bold = True
            p.add_run(pc.conflict_description)
            p2 = doc.add_paragraph()
            p2.add_run("Suggested language: ").bold = True
            p2.add_run(pc.suggested_alternative_language)


def _render_s9_5(doc: Any, findings: Any) -> None:
    """S9-5 — Agreement Summary: plain-language first, risk table, attorney questions."""
    doc.add_heading("Plain-Language Summary", 1)
    doc.add_paragraph("Written for a non-attorney commissioner — 2-minute read.").italic = True
    doc.add_paragraph(findings.plain_language_summary)

    if findings.risk_flags:
        doc.add_heading("Risk Flags", 1)
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        _tbl_header(tbl, "Section / Term", "Concern", "Severity")
        for flag in findings.risk_flags:
            row = tbl.add_row().cells
            row[0].text = flag.term_or_section
            row[1].text = flag.concern
            row[2].text = flag.severity.upper()

    if findings.attorney_questions:
        doc.add_heading("Questions for Your Attorney", 1)
        for q in findings.attorney_questions:
            doc.add_heading(q.question, 3)
            p = doc.add_paragraph()
            p.add_run("Why it matters: ").bold = True
            p.add_run(q.why_important)

    if findings.one_sided_terms:
        doc.add_heading("One-Sided Terms", 1)
        for term in findings.one_sided_terms:
            doc.add_paragraph(term, style="List Bullet")

    if findings.complexity_factors_covered:
        doc.add_heading("Complexity Factors Coverage", 1)
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        _tbl_header(tbl, "Factor", "Status", "Notes")
        for cf in findings.complexity_factors_covered:
            row = tbl.add_row().cells
            row[0].text = cf.factor
            row[1].text = cf.status.replace("_", " ").title()
            row[2].text = cf.notes or ""


def _render_s10_1(doc: Any, findings: Any) -> None:
    """S10-1 — Project Case Study: five distinct published assets, each on its own page."""
    doc.add_heading("Asset 1 of 5 — Leave-Behind Case Study", 1)
    doc.add_paragraph(
        "For city managers, stakeholders, and prospect meetings. ~400–500 words.",
    ).italic = True
    doc.add_paragraph(findings.leave_behind_case_study)

    doc.add_page_break()
    doc.add_heading("Asset 2 of 5 — Website Highlight", 1)
    doc.add_paragraph("For chawq.org — third person, human story first. ~300 words.").italic = True
    doc.add_paragraph(findings.website_highlight)

    doc.add_page_break()
    doc.add_heading("Asset 3 of 5 — Press Pitch", 1)
    doc.add_paragraph(
        "For environmental and water industry editors. 2 paragraphs.",
    ).italic = True
    doc.add_paragraph(findings.press_pitch)

    doc.add_page_break()
    doc.add_heading("Asset 4 of 5 — Pull Quotes", 1)
    doc.add_paragraph(
        "Three headline-worthy sentences for design, slides, and social use.",
    ).italic = True
    for i, quote in enumerate(findings.pull_quotes, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}.  “{quote}”")

    doc.add_page_break()
    doc.add_heading("Asset 5 of 5 — Before / After Narrative", 1)
    doc.add_paragraph(
        "For general public — event programs, newsletters, community media. ~150–250 words.",
    ).italic = True
    doc.add_paragraph(findings.before_after_narrative)


def _render_lobby_1(doc: Any, findings: Any) -> None:
    """LOBBY-1 — Lobbyist Registration Check: YES/NO answer first, details below."""
    doc.add_heading("Registration Required", 1)
    if findings.registration_required:
        answer = "YES — Registration required before any lobbying activity."
    else:
        answer = "NO — No registration required for this jurisdiction."
    p = doc.add_paragraph()
    p.add_run(answer).bold = True

    if findings.timing_requirement:
        p = doc.add_paragraph()
        p.add_run("Timing: ").bold = True
        p.add_run(findings.timing_requirement)

    if findings.registration_form_url:
        p = doc.add_paragraph()
        p.add_run("Registration Form: ").bold = True
        p.add_run(str(findings.registration_form_url))

    if findings.fees:
        p = doc.add_paragraph()
        p.add_run("Fees: ").bold = True
        p.add_run(findings.fees)

    if findings.submission_office:
        p = doc.add_paragraph()
        p.add_run("Submit To: ").bold = True
        p.add_run(findings.submission_office)

    if findings.submission_method:
        p = doc.add_paragraph()
        p.add_run("Method: ").bold = True
        p.add_run(findings.submission_method)

    if findings.county_covers_municipalities is not None:
        p = doc.add_paragraph()
        p.add_run("County Registration Covers Municipalities: ").bold = True
        p.add_run("Yes" if findings.county_covers_municipalities else "No")

    if findings.reporting_requirements:
        doc.add_heading("Reporting Requirements", 1)
        for req in findings.reporting_requirements:
            doc.add_paragraph(req, style="List Bullet")


def _render_pw_1(doc: Any, findings: Any) -> None:
    """PW-1 — Conference Attendee Research: scannable table first, detail sections below."""
    doc.add_paragraph(f"Conference: {findings.conference_name}").italic = True

    doc.add_heading("Priority Contacts", 1)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    _tbl_header(tbl, "Name", "Title / Organization", "Priority", "Suggested Opener")
    for contact in findings.top_priority_contacts:
        row = tbl.add_row().cells
        row[0].text = contact.name
        row[1].text = f"{contact.title}\n{contact.organization}"
        row[2].text = contact.relevance_to_chawq.upper()
        row[3].text = contact.suggested_opener

    doc.add_heading("Why Each Contact Matters", 1)
    for contact in findings.top_priority_contacts:
        doc.add_heading(f"{contact.name} — {contact.organization}", 2)
        p = doc.add_paragraph()
        p.add_run("Title: ").bold = True
        p.add_run(contact.title)
        doc.add_paragraph(contact.why_priority)


def _render_s4_letter(doc: Any, findings: Any) -> None:
    """S4-LETTER — Champion Briefing Letter: ready-to-send letter + internal framing."""
    p = doc.add_paragraph()
    p.add_run("Subject: ").bold = True
    p.add_run(findings.subject_line)

    doc.add_heading("Briefing Letter", 1)
    doc.add_paragraph(
        "Ready to send — edit the signature block before sending.",
    ).italic = True
    doc.add_paragraph(findings.briefing_letter)

    doc.add_heading("Key Project Framing", 1)
    doc.add_paragraph("Internal use — how C-HAWQ will position this project to funders.").italic = True
    doc.add_paragraph(findings.key_project_framing)

    doc.add_heading("Agreed Next Steps", 1)
    for step in findings.agreed_next_steps:
        doc.add_paragraph(step, style="List Bullet")


def _render_s7_plan(doc: Any, findings: Any) -> None:
    """S7-PLAN — Community Event Plan: tables for partners, run-of-show, and volunteer roles."""
    doc.add_heading("Event Overview", 1)
    doc.add_paragraph(findings.event_overview)

    if findings.community_partners:
        doc.add_heading("Community Partners", 1)
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        _tbl_header(tbl, "Organization", "Type", "Role in Event", "Contact Notes")
        for partner in findings.community_partners:
            row = tbl.add_row().cells
            row[0].text = partner.name
            row[1].text = partner.partner_type
            row[2].text = partner.role_in_event
            row[3].text = partner.contact_notes or "—"

    if findings.volunteer_framework:
        doc.add_heading("Volunteer Framework", 1)
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        _tbl_header(tbl, "Role", "# Needed", "Responsibilities")
        for vol in findings.volunteer_framework:
            row = tbl.add_row().cells
            row[0].text = vol.role
            row[1].text = str(vol.count_needed)
            row[2].text = "\n".join(f"• {r}" for r in vol.responsibilities)

    if findings.event_run_of_show:
        doc.add_heading("Run of Show", 1)
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        _tbl_header(tbl, "Time", "Activity", "Lead", "Materials")
        for seg in findings.event_run_of_show:
            row = tbl.add_row().cells
            row[0].text = seg.time_slot
            row[1].text = seg.activity
            row[2].text = seg.lead.upper()
            row[3].text = ", ".join(seg.materials_needed) if seg.materials_needed else "—"

    if findings.outreach_channels:
        doc.add_heading("Outreach Channels", 1)
        for ch in findings.outreach_channels:
            doc.add_paragraph(ch, style="List Bullet")

    if findings.permits_or_approvals_needed:
        doc.add_heading("Permits & Approvals Needed", 1)
        for item in findings.permits_or_approvals_needed:
            doc.add_paragraph(item, style="List Bullet")

    if findings.success_metrics:
        doc.add_heading("Success Metrics", 1)
        for m in findings.success_metrics:
            doc.add_paragraph(m, style="List Bullet")


def _render_s6_1(doc: Any, findings: Any) -> None:
    """S6-1 — Grant Opportunity Research: one section per grant + risks and contractors tables."""
    doc.add_heading("Grant Opportunities", 1)
    for grant in findings.grants:
        doc.add_heading(grant.name, 2)

        p = doc.add_paragraph()
        p.add_run("Agency: ").bold = True
        p.add_run(grant.administering_agency)

        if grant.typical_award_usd_min or grant.typical_award_usd_max:
            p2 = doc.add_paragraph()
            p2.add_run("Typical Award: ").bold = True
            if grant.typical_award_usd_min and grant.typical_award_usd_max:
                p2.add_run(f"${grant.typical_award_usd_min:,} – ${grant.typical_award_usd_max:,}")
            elif grant.typical_award_usd_max:
                p2.add_run(f"Up to ${grant.typical_award_usd_max:,}")
            else:
                p2.add_run(f"From ${grant.typical_award_usd_min:,}")

        p3 = doc.add_paragraph()
        p3.add_run("P3 Compatible: ").bold = True
        p3.add_run(grant.p3_compatible.replace("_", " ").title())

        p4 = doc.add_paragraph()
        p4.add_run("Deadline / Cycle: ").bold = True
        p4.add_run(grant.deadline_or_cycle)

        doc.add_heading("Eligibility", 3)
        doc.add_paragraph(grant.eligibility_summary)

        if grant.documentation_required:
            doc.add_heading("Documentation Required", 3)
            for d in grant.documentation_required:
                doc.add_paragraph(d, style="List Bullet")

        if grant.florida_precedents:
            doc.add_heading("Florida Precedents", 3)
            for fp in grant.florida_precedents:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"{fp.municipality or 'FL'} ({fp.year}): ").bold = True
                p.add_run(fp.outcome)

    if findings.risks_and_disqualifiers:
        doc.add_heading("Risks & Disqualifiers", 1)
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        _tbl_header(tbl, "Risk", "Severity", "Mitigation")
        for risk in findings.risks_and_disqualifiers:
            row = tbl.add_row().cells
            row[0].text = risk.description
            row[1].text = risk.severity.upper()
            row[2].text = risk.mitigation or "—"

    if findings.p3_contractors:
        doc.add_heading("P3 Contractors with Florida Track Record", 1)
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        _tbl_header(tbl, "Firm", "Project", "Municipality", "Outcome")
        for gc in findings.p3_contractors:
            row = tbl.add_row().cells
            row[0].text = gc.firm_name
            row[1].text = gc.project_executed
            row[2].text = gc.municipality or "—"
            row[3].text = gc.outcome


def _render_s6_3(doc: Any, findings: Any) -> None:
    """S6-3 — Commission Presentation Prep: script + objections + what-not-to-say."""
    doc.add_heading("Opening Script", 1)
    doc.add_paragraph("Verbatim 5-minute opening — read as written.").italic = True
    doc.add_paragraph(findings.opening_script)

    doc.add_heading("Top Objections", 1)
    for obj in findings.top_objections:
        label = obj.objection
        if obj.commissioner_type_likely:
            label += f"  [{obj.commissioner_type_likely}]"
        doc.add_heading(label, 3)
        doc.add_paragraph(obj.response)

    doc.add_heading("What Not to Say", 1)
    for item in findings.things_not_to_say:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Closing Statement", 1)
    doc.add_paragraph(findings.closing_statement)

    doc.add_heading("Why C-HAWQ?", 1)
    doc.add_paragraph(
        'Prepared response to “Why C-HAWQ and not a normal engineering firm?”',
    ).italic = True
    doc.add_paragraph(findings.why_chawq_answer)


def _render_s8_1(doc: Any, findings: Any) -> None:
    """S8-1 — Political Landscape Mapping: commissioner table + action plan + derailer tables."""
    doc.add_heading("Commissioner Profiles", 1)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    _tbl_header(tbl, "Commissioner", "Seat", "Risk", "Most Likely Objection")
    for cp in findings.commissioner_profiles:
        row = tbl.add_row().cells
        row[0].text = cp.name
        row[1].text = cp.seat or "—"
        row[2].text = cp.risk_level.upper()
        row[3].text = cp.most_likely_objection

    doc.add_heading("Counter-Strategies", 1)
    for cp in findings.commissioner_profiles:
        doc.add_heading(cp.name, 3)
        p = doc.add_paragraph()
        p.add_run("Most likely objection: ").bold = True
        p.add_run(cp.most_likely_objection)
        p2 = doc.add_paragraph()
        p2.add_run("Most effective counter: ").bold = True
        p2.add_run(cp.most_effective_counter)

    doc.add_heading("3-Week Action Plan", 1)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    _tbl_header(tbl, "Week", "Owner", "Task")
    for action in findings.three_week_action_plan:
        row = tbl.add_row().cells
        row[0].text = f"Week {action.week}"
        row[1].text = action.owner.upper()
        row[2].text = action.task

    if findings.high_leverage_community_voices:
        doc.add_heading("High-Leverage Community Voices", 1)
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        _tbl_header(tbl, "Name / Organization", "Leverage", "Suggested Action")
        for voice in findings.high_leverage_community_voices:
            row = tbl.add_row().cells
            row[0].text = voice.name_or_org
            row[1].text = voice.leverage_summary
            row[2].text = voice.suggested_action

    if findings.top_derailers_and_mitigations:
        doc.add_heading("Top Derailers & Mitigations", 1)
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        _tbl_header(tbl, "Derailer", "Severity", "Mitigation")
        for risk in findings.top_derailers_and_mitigations:
            row = tbl.add_row().cells
            row[0].text = risk.description
            row[1].text = risk.severity.upper()
            row[2].text = risk.mitigation or "—"


def _render_s9_1(doc: Any, findings: Any) -> None:
    """S9-1 — Project Kickoff Deck: slide cards with speaker notes + open items."""
    doc.add_heading("Kickoff Deck", 1)
    for slide in findings.slides:
        doc.add_heading(f"Slide {slide.slide_number}: {slide.section_title}", 2)
        for bp in slide.bullet_points:
            doc.add_paragraph(bp, style="List Bullet")
        if slide.speaker_notes:
            p = doc.add_paragraph()
            p.add_run("Speaker notes: ").bold = True
            p.add_run(slide.speaker_notes)

    if findings.open_items_register:
        doc.add_heading("Open Items Register", 1)
        doc.add_paragraph(
            "Assign an owner and deadline to each item at the kickoff meeting.",
        ).italic = True
        for item in findings.open_items_register:
            doc.add_paragraph(item, style="List Bullet")


def _render_s9_2(doc: Any, findings: Any) -> None:
    """S9-2 — Media & Reporter Research: combined outlets table + reporters + pitch emails."""
    all_outlets = (
        [("Local", o) for o in findings.local_outlets]
        + [("Regional Environmental", o) for o in findings.regional_environmental_press]
        + [("Municipal Trade", o) for o in findings.municipal_trade_press]
    )

    if all_outlets:
        doc.add_heading("Media Outlets", 1)
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        _tbl_header(tbl, "Outlet", "Type", "Coverage Area", "Relevant Beat")
        for category, outlet in all_outlets:
            row = tbl.add_row().cells
            row[0].text = outlet.outlet_name
            row[1].text = category
            row[2].text = outlet.coverage_area
            row[3].text = outlet.relevant_beat or "—"

    if findings.beat_reporters:
        doc.add_heading("Beat Reporters", 1)
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        _tbl_header(tbl, "Reporter", "Outlet", "Beat", "Recent Byline")
        for reporter in findings.beat_reporters:
            row = tbl.add_row().cells
            row[0].text = reporter.name
            row[1].text = reporter.outlet
            row[2].text = reporter.beat
            row[3].text = reporter.recent_relevant_byline or "—"

    if findings.wire_syndication_assessment:
        doc.add_heading("Wire Syndication Assessment", 1)
        doc.add_paragraph(findings.wire_syndication_assessment)

    if findings.pitch_emails:
        doc.add_heading("Pitch Emails", 1)
        for email in findings.pitch_emails:
            doc.add_heading(email.target_outlet_or_reporter, 2)
            p = doc.add_paragraph()
            p.add_run("Subject: ").bold = True
            p.add_run(email.subject)
            doc.add_paragraph(email.body)


_CUSTOM_RENDERERS: dict[str, Any] = {
    "LOBBY-1":   _render_lobby_1,
    "PW-1":      _render_pw_1,
    "S4-LETTER": _render_s4_letter,
    "S6-1":      _render_s6_1,
    "S6-3":      _render_s6_3,
    "S7-PLAN":   _render_s7_plan,
    "S8-1":      _render_s8_1,
    "S8-2":      _render_s8_2,
    "S8-3":      _render_s8_3,
    "S9-1":      _render_s9_1,
    "S9-2":      _render_s9_2,
    "S9-3":      _render_s9_3,
    "S9-4":      _render_s9_4,
    "S9-5":      _render_s9_5,
    "S10-1":     _render_s10_1,
}


def render_docx(brief: ResearchBrief) -> bytes:
    """Render a ResearchBrief as a Word document and return the raw bytes."""
    from docx import Document

    doc = Document()

    # ---- Title ---------------------------------------------------------------
    step_name = _TYPE_NAMES.get(brief.research_type_id, brief.research_type_id)
    doc.add_heading(
        f"{step_name} — {brief.municipality_name or 'Brief'}", 0
    )

    # ---- Metadata ------------------------------------------------------------
    meta = doc.add_paragraph()
    meta.add_run("Generated: ").bold = True
    meta.add_run(brief.generated_at.strftime("%B %d, %Y at %H:%M UTC"))
    meta.add_run("     Confidence: ").bold = True
    meta.add_run(f"{brief.overall_confidence:.2f}")
    meta.add_run("     Run: ").bold = True
    meta.add_run(brief.run_id[:8])

    # ---- Findings — type-specific or generic --------------------------------
    custom = _CUSTOM_RENDERERS.get(brief.research_type_id)
    if custom:
        custom(doc, brief.findings)
    else:
        _render_model(doc, brief.findings, level=1, skip_fields=_SKIP_ALWAYS)

    # ---- Agent notes ---------------------------------------------------------
    if brief.notes:
        doc.add_heading("Notes", 1)
        doc.add_paragraph(brief.notes)

    # ---- Sources -------------------------------------------------------------
    if brief.sources_consulted:
        doc.add_heading("Sources", 1)
        for s in brief.sources_consulted:
            doc.add_paragraph(str(s.url), style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Drive upload
# ---------------------------------------------------------------------------

def _get_drive_service():
    try:
        from googleapiclient.discovery import build
    except ImportError:
        print("ERROR: pip install google-api-python-client google-auth")
        sys.exit(1)

    access_token = os.environ.get("DRIVE_ACCESS_TOKEN")
    sa_email = os.environ.get("DRIVE_SA_EMAIL")
    sa_key = os.environ.get("DRIVE_SA_KEY")

    if access_token:
        import google.auth.credentials

        class _StaticToken(google.auth.credentials.Credentials):
            def __init__(self, token):
                super().__init__()
                self.token = token

            @property
            def valid(self):
                return True

            @property
            def expired(self):
                return False

            def refresh(self, request):
                pass

            def before_request(self, request, method, url, headers):
                self.apply(headers)

        credentials = _StaticToken(access_token)
    elif sa_email:
        from google.auth import default, impersonated_credentials
        source, _ = default()
        credentials = impersonated_credentials.Credentials(
            source_credentials=source,
            target_principal=sa_email,
            target_scopes=["https://www.googleapis.com/auth/drive"],
            lifetime=3600,
        )
    elif sa_key:
        from google.oauth2 import service_account
        credentials = service_account.Credentials.from_service_account_file(
            sa_key, scopes=["https://www.googleapis.com/auth/drive"]
        )
    else:
        from google.auth import default
        credentials, _ = default(scopes=["https://www.googleapis.com/auth/drive"])

    return build("drive", "v3", credentials=credentials, cache_discovery=False)


def _upload_or_replace(
    service, filename: str, content: bytes, mime_type: str, folder_id: str
) -> dict:
    from googleapiclient.http import MediaIoBaseUpload

    media = MediaIoBaseUpload(
        io.BytesIO(content), mimetype=mime_type, resumable=False
    )
    existing = service.files().list(
        q=f"name = '{filename}' and '{folder_id}' in parents and trashed = false",
        fields="files(id, name)",
        supportsAllDrives=True,
        includeItemsFromAllDrives=True,
    ).execute().get("files", [])

    if existing:
        return service.files().update(
            fileId=existing[0]["id"],
            media_body=media,
            fields="id, name, webViewLink, modifiedTime",
            supportsAllDrives=True,
        ).execute()

    return service.files().create(
        body={"name": filename, "parents": [folder_id]},
        media_body=media,
        fields="id, name, webViewLink, createdTime",
        supportsAllDrives=True,
    ).execute()


def upload_brief(
    brief: ResearchBrief,
    folder_id: str = DEFAULT_FOLDER_ID,
) -> dict[str, dict]:
    """Upload brief as JSON + Word doc to Drive. Returns file metadata per format."""
    service = _get_drive_service()

    json_bytes = brief.model_dump_json(indent=2).encode("utf-8")
    docx_bytes = render_docx(brief)

    return {
        "json": _upload_or_replace(
            service, filename_for(brief, "json"),
            json_bytes, "application/json", folder_id,
        ),
        "docx": _upload_or_replace(
            service, filename_for(brief, "docx"),
            docx_bytes, DOCX_MIME, folder_id,
        ),
    }
