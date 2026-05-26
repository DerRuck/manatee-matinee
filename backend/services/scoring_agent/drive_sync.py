"""Upload ScoringResults to Google Drive as Word documents.

Auth precedence: DRIVE_SA_EMAIL (impersonate via ADC) > DRIVE_SA_KEY (file) > ADC.

Usage:
    from services.scoring_agent.drive_sync import upload_score
    result = upload_score(scoring_result, folder_id="1L-zcN4jA83EfsrRyei_ewbKOEKMKz-lC")
    # result = {"docx": {...Drive file metadata...}, "json": {...}}

Files land in {folder_id}/{contact folder}/Pipeline Scores/.
Contact subfolder is found or created on first use (idempotent).
Re-runs overwrite files with the same name so the folder doesn't accumulate stale copies.
"""
from __future__ import annotations

import io
import re
from typing import Any

from services.scoring_agent.schema import ScoringResult

DEFAULT_FOLDER_ID = "1L-zcN4jA83EfsrRyei_ewbKOEKMKz-lC"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

_HEAT_LABEL = {
    "boil":   "BOIL",
    "simmer": "SIMMER",
    "stall":  "STALL",
    "cold":   "COLD",
    "won":    "WON",
    "lost":   "LOST",
}

_IMPACT_PREFIX = {
    "positive": "[+]",
    "negative": "[-]",
    "neutral":  "[ ]",
}


# ---------------------------------------------------------------------------
# Display label + filename
# ---------------------------------------------------------------------------

def display_label(result: ScoringResult) -> str:
    """Readable identifier for one scored contact.

    Priority chain — pick the first non-empty:
      1. contact_name        (e.g. 'Jamie Sheehan')
      2. contact_email
      3. municipality_name
      4. contact_id          (last resort — the opaque GHL id)

    Used by both the filename and the docx header so the team never has to
    eyeball an id to figure out which contact a file belongs to.
    """
    candidates = (
        result.contact_name,
        result.contact_email,
        result.municipality_name,
        result.contact_id,
    )
    for value in candidates:
        if value:
            cleaned = str(value).strip()
            if cleaned:
                return cleaned
    return "Contact"


_FILENAME_SAFE_RE = re.compile(r"[^A-Za-z0-9._-]+")
_COLLAPSE_UNDERSCORES_RE = re.compile(r"_+")


def _slug_for_filename(value: str) -> str:
    """Lowercase, ASCII-safe slug — emails keep their structure via `_at_`."""
    s = value.replace("@", "_at_").replace(" ", "_")
    s = _FILENAME_SAFE_RE.sub("_", s)
    s = _COLLAPSE_UNDERSCORES_RE.sub("_", s)
    s = s.strip("._-").lower()
    return s or "contact"


def filename_for(result: ScoringResult, ext: str) -> str:
    """`pipeline_score_<readable-slug>_<run-prefix>.<ext>`."""
    slug = _slug_for_filename(display_label(result))
    return f"pipeline_score_{slug}_{result.run_id[:8]}.{ext}"


# ---------------------------------------------------------------------------
# Docx renderer
# ---------------------------------------------------------------------------

def render_docx(result: ScoringResult) -> bytes:
    """Render a ScoringResult as a branded Word document."""
    from docx import Document
    from services.branding.docx_styles import (
        add_brand_header, add_meta_line, apply_brand_styles,
    )

    doc = Document()
    apply_brand_styles(doc)

    f = result.findings
    primary_label = display_label(result)
    # When the primary label is a person's name (or email), surface the
    # municipality as secondary context so the doc header still tells you
    # which city this lead is in.
    secondary = (
        result.municipality_name
        if result.municipality_name and result.municipality_name != primary_label
        else None
    )
    subtitle_bits = [f"PIPELINE-SCORE v{result.prompt_version} · {result.triggered_by}"]
    if secondary:
        subtitle_bits.append(secondary)

    add_brand_header(
        doc,
        title=f"Pipeline Score — {primary_label}",
        subtitle=" · ".join(subtitle_bits),
    )

    add_meta_line(
        doc,
        generated=result.generated_at.strftime("%B %d, %Y at %H:%M UTC"),
        run=result.run_id[:8],
    )

    # ---- Score summary -------------------------------------------------------
    doc.add_heading("Score Summary", 1)

    heat_label = _HEAT_LABEL.get(f.lead_heat, f.lead_heat.upper())
    confidence_pct = round(f.step_confidence * 100)

    doc.add_paragraph(
        f"{f.current_step_name}  (Phase {f.current_phase})"
    )
    doc.add_paragraph(
        f"Heat: {heat_label}   Score: {f.lead_heat_score}/100   "
        f"Step confidence: {confidence_pct}%"
    )

    if f.days_since_last_signal is not None:
        doc.add_paragraph(f"Days since last signal: {f.days_since_last_signal}")

    # Summary one-liner in its own styled paragraph
    p = doc.add_paragraph()
    run = p.add_run(f.summary_one_line)
    run.bold = True

    # ---- Advancement status --------------------------------------------------
    doc.add_heading("Advancement Status", 1)

    status = "Ready to advance" if f.ready_to_advance else "Not ready to advance"
    doc.add_paragraph(status)

    if f.next_step_blockers:
        doc.add_heading("Blockers", 2)
        for blocker in f.next_step_blockers:
            owner_tag = f" ({blocker.owner})" if blocker.owner else ""
            doc.add_paragraph(
                f"[{blocker.severity.upper()}]{owner_tag} {blocker.description}",
                style="List Bullet",
            )

    # ---- Signals -------------------------------------------------------------
    if f.signals:
        doc.add_heading("Signals", 1)
        for sig in f.signals:
            prefix = _IMPACT_PREFIX.get(sig.impact, "[ ]")
            weight_tag = f"  (w={sig.weight:.2f})"
            doc.add_paragraph(
                f"{prefix}{weight_tag} {sig.description}",
                style="List Bullet",
            )
            sub = doc.add_paragraph(f"    Source: {sig.evidence_source}")
            sub.paragraph_format.left_indent = _inches(0.25)
            _set_small_font(sub)

    # ---- Recommended actions -------------------------------------------------
    if f.recommended_actions:
        doc.add_heading("Recommended Actions", 1)
        for action in f.recommended_actions:
            step_tag = f"[Step {action.proven_process_step}] " if action.proven_process_step else ""
            due = f"≤{action.due_within_days}d" if action.due_within_days > 0 else "today"
            doc.add_paragraph(
                f"({action.owner}, {due}) {step_tag}{action.action}",
                style="List Bullet",
            )

    # ---- Go/No-Go scorecard (Step 3+) ----------------------------------------
    if f.go_no_go is not None:
        gng = f.go_no_go
        doc.add_heading("Go / No-Go Scorecard", 1)
        doc.add_paragraph(f"Decision: {gng.decision}   Total: {gng.total}/21")
        doc.add_paragraph(gng.rationale)

        scorecard_rows = [
            ("Authority",           gng.authority_score),
            ("Project specificity", gng.project_specificity_score),
            ("Solvability",         gng.solvability_score),
            ("Champion passion",    gng.champion_passion_score),
            ("C-HAWQ fit",          gng.chawq_fit_score),
            ("P3 candidacy",        gng.p3_candidacy_score),
            ("Political readiness", gng.political_readiness_score),
        ]
        for label, score in scorecard_rows:
            doc.add_paragraph(f"{label}: {score}/3", style="List Bullet")

    # ---- Boil criteria -------------------------------------------------------
    if f.boil_criteria:
        doc.add_heading("Boil Criteria", 1)
        for criterion in f.boil_criteria:
            evidence_tag = f" — {criterion.evidence}" if criterion.evidence else ""
            doc.add_paragraph(
                f"{criterion.key}: {criterion.answer.upper()}{evidence_tag}",
                style="List Bullet",
            )

    # ---- Agent notes ---------------------------------------------------------
    if result.notes:
        doc.add_heading("Notes", 1)
        doc.add_paragraph(result.notes)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Drive upload
# ---------------------------------------------------------------------------

def upload_score(
    result: ScoringResult,
    folder_id: str = DEFAULT_FOLDER_ID,
) -> dict[str, dict]:
    """Upload score as JSON + Word doc to Drive. Returns file metadata per format.

    Files land in {folder_id}/{contact folder}/Pipeline Scores/.
    Contact and type subfolders are created on first use and reused on re-runs.
    """
    from services.drive.folders import ensure_subfolder, resolve_contact_folder_name

    service = _get_drive_service()

    contact_folder_id = ensure_subfolder(
        service, folder_id, resolve_contact_folder_name(result),
    )
    target_folder_id = ensure_subfolder(
        service, contact_folder_id, "Pipeline Scores",
    )

    json_bytes = result.model_dump_json(indent=2).encode("utf-8")
    docx_bytes = render_docx(result)

    return {
        "json": _upload_or_replace(
            service, filename_for(result, "json"),
            json_bytes, "application/json", target_folder_id,
        ),
        "docx": _upload_or_replace(
            service, filename_for(result, "docx"),
            docx_bytes, DOCX_MIME, target_folder_id,
        ),
    }


# ---------------------------------------------------------------------------
# Shared Drive helpers (mirrors research_agent/drive_sync.py)
# ---------------------------------------------------------------------------

def _get_drive_service() -> Any:
    from googleapiclient.discovery import build

    sa_email = _drive_sa_email()
    sa_key = _drive_sa_key()

    if sa_email:
        from google.auth import default, impersonated_credentials
        source_creds, _ = default(scopes=["https://www.googleapis.com/auth/drive"])
        credentials = impersonated_credentials.Credentials(
            source_credentials=source_creds,
            target_principal=sa_email,
            target_scopes=["https://www.googleapis.com/auth/drive"],
        )
    elif sa_key:
        from google.oauth2 import service_account
        credentials = service_account.Credentials.from_service_account_file(
            sa_key,
            scopes=["https://www.googleapis.com/auth/drive"],
        )
    else:
        from google.auth import default
        credentials, _ = default(scopes=["https://www.googleapis.com/auth/drive"])

    return build("drive", "v3", credentials=credentials, cache_discovery=False)


def _drive_sa_email() -> str | None:
    try:
        from core.settings import get_settings
        return get_settings().drive_sa_email or None
    except Exception:
        import os
        return os.environ.get("DRIVE_SA_EMAIL") or None


def _drive_sa_key() -> str | None:
    try:
        from core.settings import get_settings
        return get_settings().drive_sa_key_path or None
    except Exception:
        import os
        return os.environ.get("DRIVE_SA_KEY_PATH") or None


def _upload_or_replace(
    service: Any, filename: str, content: bytes, mime_type: str, folder_id: str,
) -> dict:
    from googleapiclient.http import MediaIoBaseUpload

    media = MediaIoBaseUpload(
        io.BytesIO(content), mimetype=mime_type, resumable=False,
    )
    escaped = filename.replace("\\", "\\\\").replace("'", "\\'")
    existing = service.files().list(
        q=f"name = '{escaped}' and '{folder_id}' in parents and trashed = false",
        fields="files(id, name)",
        supportsAllDrives=True,
        includeItemsFromAllDrives=True,
    ).execute().get("files", [])

    if existing:
        return service.files().update(
            fileId=existing[0]["id"],
            media_body=media,
            fields="id, name, webViewLink, modifiedTime",
            supportsAllDrives=True,
        ).execute()

    return service.files().create(
        body={"name": filename, "parents": [folder_id]},
        media_body=media,
        fields="id, name, webViewLink, createdTime",
        supportsAllDrives=True,
    ).execute()


# ---------------------------------------------------------------------------
# Docx helpers
# ---------------------------------------------------------------------------

def _inches(n: float) -> Any:
    from docx.shared import Inches
    return Inches(n)


def _set_small_font(paragraph: Any) -> None:
    from docx.shared import Pt
    for run in paragraph.runs:
        run.font.size = Pt(9)
