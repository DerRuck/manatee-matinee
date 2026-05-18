"""
Google Drive client.

V1 responsibilities:
  - Watch a folder (push notification channel -> /webhooks/drive).
  - Read files in place (never moved, per Guardrail #1).
  - Mirror agent outputs to /CHAWQ/<city>/<project>/<stage>/... folders.

Sprint demo: just upload_text_file() so the Hello World agent's output can
land in the watched Drive folder. Watch-channel + read paths come in later Sprint.

Auth strategy:
  - In Cloud Run (K_SERVICE env set): default() returns the runtime SA's
    creds via the metadata server. We just request the Drive scope.
  - Local dev: default() returns user OAuth creds, which can't directly hit
    Drive (this is the "This app is blocked" wall we hit in the watch spike).
    So we impersonate the runtime SA — same pattern as scripts/drive_watch.py.
"""
from __future__ import annotations

import io
import logging
import os
from typing import Any, Iterator

from google.auth import default, impersonated_credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload, MediaIoBaseUpload

from core.settings import Settings, get_settings

logger = logging.getLogger(__name__)


SA_EMAIL = "chawq-api-runtime@chawq-manatee-matinee.iam.gserviceaccount.com"
DRIVE_SCOPE = "https://www.googleapis.com/auth/drive"

GOOGLE_DOC_MIME = "application/vnd.google-apps.document"

# Drive folders carry their own mime type. Used by walk_folder() to decide
# what to recurse into vs. yield as a leaf.
FOLDER_MIME = "application/vnd.google-apps.folder"

# Plain text-ish mime types we can download as raw bytes.
PLAIN_TEXT_MIMES = {
    "text/plain",
    "text/markdown",
    "text/x-markdown",
    "application/octet-stream",  # Drive sometimes mislabels .md as octet-stream
}

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

PDF_MIME = "application/pdf"

# Mime prefixes / exact types we never try to extract text from. Audio
# and images need transcription / OCR (out of V1 scope); JSON dumps from
# Plaud are noisy raw data and shouldn't be embed targets. Caller can
# check is_text_extractable_mime() to decide before download.
SKIP_MIME_PREFIXES: tuple[str, ...] = (
    "audio/",
    "video/",
    "image/",
)
SKIP_MIME_EXACT: frozenset[str] = frozenset({
    "application/json",
})


def is_text_extractable_mime(mime_type: str | None) -> bool:
    """True if download_file_as_text() will return text for this mime type."""
    if not mime_type:
        return False
    if mime_type.startswith(SKIP_MIME_PREFIXES):
        return False
    if mime_type in SKIP_MIME_EXACT:
        return False
    return mime_type in PLAIN_TEXT_MIMES or mime_type in {
        GOOGLE_DOC_MIME,
        DOCX_MIME,
        PDF_MIME,
    }


_drive_service: Any | None = None


def _get_drive_service():
    """Lazy-build (and cache) the Drive v3 client."""
    global _drive_service
    if _drive_service is not None:
        return _drive_service

    if os.environ.get("K_SERVICE"):
        creds, _ = default(scopes=[DRIVE_SCOPE])
    else:
        source_creds, _ = default()
        creds = impersonated_credentials.Credentials(
            source_credentials=source_creds,
            target_principal=SA_EMAIL,
            target_scopes=[DRIVE_SCOPE],
            lifetime=3600,
        )

    _drive_service = build("drive", "v3", credentials=creds, cache_discovery=False)
    return _drive_service


def find_or_create_folder(name: str, parent_folder_id: str) -> str:
    """
    Return the Drive folder ID for `name` under `parent_folder_id`.
    Creates the folder if it doesn't exist; returns the existing ID if
    one matches by exact name.

    Used by agent runners to materialize the per-lead Drive layout
    (`/Leads/<municipality>/<contact>/<typed_subfolder>/...`) the first
    time an agent writes for a new lead. Idempotent — re-runs are cheap.

    Raises HttpError on Drive API failure (caller decides whether to
    retry or fall back to writing flat).
    """
    service = _get_drive_service()
    # Look for an existing folder of that name under the parent.
    children = list_folder_files(parent_folder_id)
    for child in children:
        if (
            child.get("mimeType") == FOLDER_MIME
            and child.get("name") == name
        ):
            return child["id"]

    # Not found — create it.
    body = {
        "name": name,
        "mimeType": FOLDER_MIME,
        "parents": [parent_folder_id],
    }
    response = (
        service.files()
        .create(body=body, fields="id, name", supportsAllDrives=True)
        .execute()
    )
    logger.info(
        "drive folder created",
        extra={
            "folder_id": response["id"],
            "folder_name": name,
            "parent_folder_id": parent_folder_id,
        },
    )
    return response["id"]


def upload_text_file(
    folder_id: str,
    filename: str,
    content: str,
    mime_type: str = "text/markdown",
) -> dict[str, str]:
    """
    Upload `content` as a new file named `filename` inside the Drive folder
    `folder_id`. Returns the response dict containing at least:
      - id           : the Drive file ID
      - name         : the final filename
      - webViewLink  : a clickable link to the file in Drive

    Caller is responsible for ensuring the runtime SA has Editor access on
    the target folder. Without it this will fail with insufficientPermissions.
    """
    return upload_bytes_file(
        folder_id=folder_id,
        filename=filename,
        content=content.encode("utf-8"),
        mime_type=mime_type,
    )


def upload_bytes_file(
    folder_id: str,
    filename: str,
    content: bytes,
    mime_type: str,
) -> dict[str, str]:
    """
    Upload raw bytes as a new file inside `folder_id`. Used for binary
    artifacts the agents produce — PDF letters, future pptx decks, etc.
    Returns the same shape as upload_text_file: id, name, webViewLink.

    text/markdown uploads should keep using upload_text_file (this
    function is the lower-level path it now delegates to).
    """
    service = _get_drive_service()
    media = MediaIoBaseUpload(
        io.BytesIO(content),
        mimetype=mime_type,
        resumable=False,
    )
    body = {"name": filename, "parents": [folder_id]}

    response = (
        service.files()
        .create(
            body=body,
            media_body=media,
            fields="id, name, webViewLink",
            supportsAllDrives=True,
        )
        .execute()
    )
    return response


def list_folder_files(
    folder_id: str,
    page_size: int = 100,
) -> list[dict[str, Any]]:
    """
    List non-trashed children of `folder_id` (one level only — no recursion).

    Returns Drive Files responses with these fields populated:
      - id, name, mimeType, modifiedTime, webViewLink, parents

    Children include subfolders (mimeType == FOLDER_MIME); callers that want
    to flatten file trees should use `walk_folder()` instead.
    """
    service = _get_drive_service()
    files: list[dict[str, Any]] = []
    page_token: str | None = None

    while True:
        resp = (
            service.files()
            .list(
                q=f"'{folder_id}' in parents and trashed = false",
                fields=(
                    "nextPageToken, "
                    "files(id, name, mimeType, modifiedTime, webViewLink, parents, size)"
                ),
                pageSize=page_size,
                pageToken=page_token,
                supportsAllDrives=True,
                includeItemsFromAllDrives=True,
            )
            .execute()
        )
        files.extend(resp.get("files", []))
        page_token = resp.get("nextPageToken")
        if not page_token:
            break

    return files


def get_file_metadata(file_id: str) -> dict[str, Any]:
    """
    Fetch a single file's metadata.

    Returns the same field shape that `list_folder_files()` and
    `walk_folder()` yield (id, name, mimeType, modifiedTime, webViewLink,
    parents) plus `size`, so the result drops cleanly into any code path
    that already understands walker output. `size` is needed by the
    iflytek resolver to skip 0-byte sidecars.
    """
    service = _get_drive_service()
    return (
        service.files()
        .get(
            fileId=file_id,
            fields="id, name, mimeType, modifiedTime, webViewLink, parents, size",
            supportsAllDrives=True,
        )
        .execute()
    )


def derive_path_segments(
    file_id: str,
    stop_at_folder_id: str,
    max_depth: int = 16,
) -> list[str]:
    """
    Walk parents back from `file_id` up to (but not including)
    `stop_at_folder_id`, returning folder names in root → leaf order.

    Mirrors what `walk_folder()` yields as path_segments for a file at
    the same location: the file's immediate parent name is last in the
    list; `stop_at_folder_id`'s name is NOT included. Files directly
    inside `stop_at_folder_id` get an empty list back.

    Used by the single-file ingest entry point so per-source resolvers
    see the same lineage shape they get during a recursive folder walk.

    Raises ValueError if the file isn't actually a descendant of
    `stop_at_folder_id` within `max_depth` levels — keeps mis-classified
    files from sneaking into the wrong source taxonomy silently.
    """
    service = _get_drive_service()

    # Start from the file's parent, not the file itself.
    file_meta = (
        service.files()
        .get(fileId=file_id, fields="parents", supportsAllDrives=True)
        .execute()
    )
    parents = file_meta.get("parents") or []
    if not parents:
        raise ValueError(
            f"file {file_id} has no parents — can't derive path segments"
        )

    segments: list[str] = []
    current = parents[0]

    for _ in range(max_depth):
        if current == stop_at_folder_id:
            return list(reversed(segments))

        folder = (
            service.files()
            .get(
                fileId=current,
                fields="id, name, parents",
                supportsAllDrives=True,
            )
            .execute()
        )
        segments.append(folder["name"])
        next_parents = folder.get("parents") or []
        if not next_parents:
            break  # hit My Drive root without seeing stop_at_folder_id
        current = next_parents[0]

    raise ValueError(
        f"file {file_id} is not a descendant of folder "
        f"{stop_at_folder_id} within {max_depth} levels"
    )


def walk_folder(
    root_folder_id: str,
    page_size: int = 100,
    max_depth: int = 16,
) -> Iterator[tuple[dict[str, Any], list[str]]]:
    """
    Recursively walk a Drive folder tree, yielding non-folder leaves.

    Each yield is a tuple `(file_meta, path_segments)`:
      - file_meta: same shape as `list_folder_files()` entries.
      - path_segments: folder names from the root *down to* the file's
        immediate parent. Files directly inside `root_folder_id` get an
        empty list; files one level deeper get a single-element list, etc.

    Subfolders are recursed into, never yielded. Trashed items are skipped.
    `max_depth` is a safety stop against unexpectedly deep trees (the V1
    corpus convention is two levels: <municipality>/<document_type>).

    The root folder's own name is NOT included in path_segments; callers
    that want it can pass `root_folder_id` to `get_file_metadata()`.
    """
    # Breadth-first queue of (folder_id, path_so_far).
    stack: list[tuple[str, list[str], int]] = [(root_folder_id, [], 0)]

    while stack:
        folder_id, path, depth = stack.pop(0)
        if depth > max_depth:
            logger.warning(
                "walk_folder depth limit hit — skipping",
                extra={"folder_id": folder_id, "path": "/".join(path), "depth": depth},
            )
            continue

        children = list_folder_files(folder_id, page_size=page_size)
        for child in children:
            mime = child.get("mimeType")
            if mime == FOLDER_MIME:
                stack.append((child["id"], path + [child["name"]], depth + 1))
                continue
            yield child, list(path)


def download_file_as_text(
    file_id: str,
    mime_type: str,
    file_name: str | None = None,
) -> str | None:
    """
    Download a Drive file's content as plain UTF-8 text.

      - Google Doc -> export as text/plain
      - text/plain or text/markdown -> raw bytes -> utf-8
      - .docx -> raw bytes -> python-docx -> paragraphs + table cells joined
      - .pdf -> raw bytes -> pypdf -> page text concatenated (native text only)
      - audio/video/image/json -> None (skip with log)
      - Anything else -> None (caller logs and skips)

    `file_name` is used to disambiguate `application/octet-stream` files,
    which Drive serves for both `.md` text and `.opus` audio. When given,
    octet-stream files are only treated as text if the extension says so.
    Without a filename we fall back to the old behavior (treat as text).
    """
    if mime_type and (
        mime_type.startswith(SKIP_MIME_PREFIXES) or mime_type in SKIP_MIME_EXACT
    ):
        logger.info(
            "drive download skip — non-text mime type",
            extra={"file_id": file_id, "mime_type": mime_type, "file_name": file_name},
        )
        return None

    service = _get_drive_service()

    if mime_type == GOOGLE_DOC_MIME:
        result = (
            service.files()
            .export(fileId=file_id, mimeType="text/plain")
            .execute()
        )
        return result.decode("utf-8") if isinstance(result, bytes) else result

    if mime_type in PLAIN_TEXT_MIMES:
        if mime_type == "application/octet-stream" and file_name:
            ext = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""
            if ext not in {"md", "markdown", "txt"}:
                logger.info(
                    "drive download skip — octet-stream with non-text extension",
                    extra={"file_id": file_id, "file_name": file_name, "ext": ext},
                )
                return None
        return _download_bytes(service, file_id).decode("utf-8", errors="replace")

    if mime_type == DOCX_MIME:
        return _docx_bytes_to_text(_download_bytes(service, file_id))

    if mime_type == PDF_MIME:
        return _pdf_bytes_to_text(_download_bytes(service, file_id))

    logger.info(
        "drive download skip — unsupported mime type",
        extra={"file_id": file_id, "mime_type": mime_type, "file_name": file_name},
    )
    return None


def _download_bytes(service, file_id: str) -> bytes:
    """Internal: get_media -> in-memory bytes."""
    request = service.files().get_media(fileId=file_id)
    buf = io.BytesIO()
    downloader = MediaIoBaseDownload(buf, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()
    return buf.getvalue()


def _pdf_bytes_to_text(data: bytes) -> str | None:
    """
    Extract text from a PDF's raw bytes using pypdf.

    V1 handles native (text-bearing) PDFs only — exported docs from Word,
    Drive, Iflytek, Gmail-as-PDF, etc. Scanned/image-only PDFs come back
    near-empty; we return None in that case so the caller can skip the
    file rather than embed a corpus of whitespace.

    Encrypted PDFs are also returned as None — V1 doesn't unlock them.
    """
    # Lazy import: pypdf is only needed when we hit a .pdf file.
    from pypdf import PdfReader  # type: ignore
    from pypdf.errors import PdfReadError  # type: ignore

    try:
        reader = PdfReader(io.BytesIO(data))
    except PdfReadError as exc:
        logger.warning("pdf parse error — skipping", extra={"error": str(exc)})
        return None

    if reader.is_encrypted:
        logger.info("pdf is encrypted — skipping")
        return None

    parts: list[str] = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception as exc:
            # pypdf occasionally throws on malformed pages; skip the page
            # rather than failing the whole file.
            logger.warning("pdf page extract error — skipping page", extra={"error": str(exc)})
            continue
        text = text.strip()
        if text:
            parts.append(text)

    if not parts:
        # Native PDF text extraction returned nothing — likely scanned.
        # OCR is a separate decision (out of V1 scope).
        return None

    return "\n\n".join(parts)


def _docx_bytes_to_text(data: bytes) -> str:
    """
    Extract text from a .docx file's raw bytes.

    Pulls paragraph text and table cells in document order. Skips empty
    paragraphs and trims whitespace; doesn't attempt to preserve formatting,
    images, comments, or footnotes — none of those help embedding quality.
    """
    # Lazy import: python-docx is only needed when we hit a .docx file.
    from docx import Document as DocxDocument  # type: ignore

    doc = DocxDocument(io.BytesIO(data))

    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            row_text = " | ".join(c for c in cells if c)
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


class DriveClient:
    def __init__(self, settings: Settings | None = None) -> None:
        self.settings = settings or get_settings()
        self._service = None

    async def register_watch(self, folder_id: str, webhook_url: str) -> dict:
        """POST /files/{fileId}/watch — register a push notification channel."""
        raise NotImplementedError("Drive watch registration — Sprint 2 task.")

    async def mirror_file(
        self,
        source_path: str,
        target_folder_path: str,
        filename: str,
    ) -> str:
        """Upload a local file to the target Drive folder. Returns Drive file ID."""
        raise NotImplementedError("Drive mirror write — Sprint 3 one-way sync task.")